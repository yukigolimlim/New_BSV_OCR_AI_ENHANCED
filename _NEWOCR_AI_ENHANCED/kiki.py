"""
DocExtract Pro — Banco San Vicente
====================================
Radiant Navy · White · Lime  |  CustomTkinter
v4 Changes:
  • Added "Analyze Loan" button — runs LLM credit assessment on extracted text
  • Credit scoring gracefully skipped if sheet is empty/all-zero
  • Analysis runs in background thread with spinner
  • Result shown in a separate Analysis tab (right panel tabview)
  • Minimize fix: ctypes ShowWindow, no overrideredirect toggle
  • _fix_windows_taskbar properly inside DocExtractorApp class
  • Loan product catalog (banco_san_vicente_loans.json) injected into analysis prompt
  • Section 10: Loan Eligibility — eligible products, loanable amounts, upgrade path
"""
import ctypes
import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog
from tkinter import font as tkfont
import threading
import math
import sys
import os
import re
import base64
from pathlib import Path
from dotenv import load_dotenv
import json

load_dotenv()

ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")

# ══════════════════════════════════════════════════════════════════════════════
#  PALETTE
# ══════════════════════════════════════════════════════════════════════════════
NAVY_DEEP    = "#1C2E7A"
NAVY         = "#2B45A8"
NAVY_MID     = "#3858C8"
NAVY_LIGHT   = "#4E6EE0"
NAVY_PALE    = "#7A96F0"
NAVY_GHOST   = "#C8D4F8"
NAVY_MIST    = "#E8EEFF"

WHITE        = "#FFFFFF"
OFF_WHITE    = "#F5F7FF"
CARD_WHITE   = "#FFFFFF"
PANEL_LEFT   = "#F0F3FF"

LIME_BRIGHT  = "#C8F020"
LIME         = "#A8D818"
LIME_MID     = "#90C010"
LIME_DARK    = "#6A9408"
LIME_PALE    = "#E0F870"
LIME_MIST    = "#F2FFCC"

BORDER_LIGHT = "#D4DCF8"
BORDER_MID   = "#B0BEEC"

TXT_NAVY     = "#1C2E7A"
TXT_NAVY_MID = "#3858C8"
TXT_SOFT     = "#6878B8"
TXT_MUTED    = "#9AAACE"
TXT_ON_LIME  = "#1C2E7A"

ACCENT_GOLD    = "#F0A800"
ACCENT_SUCCESS = "#22C870"
ACCENT_RED     = "#E74C3C"

WIN_W = 1240
WIN_H = 780

SCRIPT_DIR   = Path(__file__).parent
LOGO_PATH    = SCRIPT_DIR / "bsv_logotxt.png"
IMAGE_EXTS   = {".png",".jpg",".jpeg",".bmp",".tiff",".tif",".webp",".gif"}
POPPLER_PATH = r"C:\poppler\Release-25.12.0-0\poppler-25.12.0\Library\bin"


# ── Colour blend ──────────────────────────────────────────────────────────────
def _hex_blend(c1, c2, t):
    r1,g1,b1 = int(c1[1:3],16),int(c1[3:5],16),int(c1[5:7],16)
    r2,g2,b2 = int(c2[1:3],16),int(c2[3:5],16),int(c2[5:7],16)
    return f"#{int(r1+(r2-r1)*t):02x}{int(g1+(g2-g1)*t):02x}{int(b1+(b2-b1)*t):02x}"


# ── Font cascade ──────────────────────────────────────────────────────────────
def _best_font():
    import tkinter.font as tkfont
    available = set(tkfont.families())
    for f in ("Nunito","Montserrat","Poppins","Segoe UI","Calibri","Arial"):
        if f in available:
            return f
    return "Arial"


def _register_fonts():
    """Load Montserrat from the project folder into Tkinter."""
    try:
        import pyglet
        font_dir = SCRIPT_DIR
        for ttf in font_dir.glob("Montserrat*.ttf"):
            pyglet.font.add_file(str(ttf))
    except ImportError:
        pass

_FONT_FAMILY = None

def F(size, weight="normal"):
    global _FONT_FAMILY
    if _FONT_FAMILY is None:
        _FONT_FAMILY = _best_font()
    return (_FONT_FAMILY, size, weight)

def FMONO(size, weight="normal"):
    import tkinter.font as tkfont
    available = set(tkfont.families())
    for f in ("JetBrains Mono","Cascadia Code","Consolas","Courier New"):
        if f in available:
            return (f, size, weight)
    return ("Courier New", size, weight)


# ══════════════════════════════════════════════════════════════════════════════
#  GROQ / PaddleOCR
# ══════════════════════════════════════════════════════════════════════════════
_groq_client = None

def _get_groq():
    global _groq_client
    if _groq_client is None:
        from groq import Groq
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            raise EnvironmentError(
                "GROQ_API_KEY not found.\n\nCreate a .env file:\n"
                "    GROQ_API_KEY=your_key_here\n\nhttps://console.groq.com"
            )
        _groq_client = Groq(api_key=api_key)
    return _groq_client

_paddle_ocr = None

def _get_paddle_ocr():
    global _paddle_ocr
    if _paddle_ocr is None:
        from paddleocr import PaddleOCR
        _paddle_ocr = PaddleOCR(use_angle_cls=True, lang="en", use_gpu=False)
    return _paddle_ocr


# ══════════════════════════════════════════════════════════════════════════════
#  LOAN CATALOG LOADER
# ══════════════════════════════════════════════════════════════════════════════
def _build_loan_catalog_text():
    """
    Read banco_san_vicente_loans.json and return a compact human-readable
    summary suitable for injection into an LLM prompt.
    """
    json_path = SCRIPT_DIR / "banco_san_vicente_loans.json"
    if not json_path.exists():
        return "[Loan product catalog file not found — skip eligibility section]"
    try:
        with open(json_path, "r", encoding="utf-8") as f:
            catalog = json.load(f)

        lines = ["=== BANCO SAN VICENTE LOAN PRODUCT CATALOG ===\n"]
        for cat in catalog.get("categories", []):
            lines.append(f"\n[{cat['category_name'].upper()}]")
            for p in cat.get("products", []):
                lines.append(f"\n  Product : {p['product_name']}  (id: {p['product_id']})")
                if p.get("purpose"):
                    lines.append(f"  Purpose : {p['purpose']}")
                if p.get("qualification"):
                    lines.append(f"  Qualify : {p['qualification']}")

                # ── Loan amount ──────────────────────────────────────────
                la = p.get("loan_amount")
                if la and isinstance(la, dict):
                    amt_parts = []
                    if la.get("minimum"):
                        amt_parts.append(f"min ₱{la['minimum']:,}")
                    if la.get("maximum") and isinstance(la["maximum"], (int, float)):
                        amt_parts.append(f"max ₱{la['maximum']:,}")
                    elif la.get("maximum"):
                        amt_parts.append(f"max: {la['maximum']}")
                    if la.get("first_loan_maximum"):
                        amt_parts.append(f"first-loan max ₱{la['first_loan_maximum']:,}")
                    if la.get("first_cycle"):
                        fc = la["first_cycle"]
                        amt_parts.append(
                            f"first cycle ₱{fc.get('min',0):,}–₱{fc.get('max',0):,}"
                        )
                    if la.get("calculation"):
                        amt_parts.append(f"calc: {la['calculation']}")
                    # franchise sub-keys (Pampasada)
                    for fk in ("daet", "labo"):
                        fv = la.get("franchise_value", {})
                        if isinstance(fv, dict) and fv.get(fk):
                            amt_parts.append(f"franchise ({fk}) ₱{fv[fk]:,}")
                    if amt_parts:
                        lines.append(f"  Amount  : {', '.join(amt_parts)}")

                # ── Interest rates ───────────────────────────────────────
                rates = p.get("interest_rates") or []
                for r in rates:
                    if isinstance(r, dict) and r.get("term") and r.get("interest_rate"):
                        lines.append(
                            f"  Rate    : {r['term']} @ {r['interest_rate']}"
                            f"  SC {r.get('service_charge','N/A')}"
                        )

                # ── Collateral ───────────────────────────────────────────
                col = p.get("collateral")
                if col:
                    if isinstance(col, list):
                        lines.append(f"  Collat  : {', '.join(col)}")
                    elif isinstance(col, str):
                        lines.append(f"  Collat  : {col}")
                    elif isinstance(col, dict):
                        for tier_k, tier_v in col.items():
                            if isinstance(tier_v, dict) and tier_v.get("range"):
                                rng   = tier_v["range"]
                                items = tier_v.get("items", [])
                                lines.append(
                                    f"  Collat ({tier_k} "
                                    f"₱{rng.get('min',0):,}–₱{rng.get('max',0):,}): "
                                    f"{', '.join(items)}"
                                )
                            elif isinstance(tier_v, list):
                                lines.append(
                                    f"  Collat ({tier_k}): {', '.join(tier_v)}"
                                )

                # ── Notes ────────────────────────────────────────────────
                if p.get("other_notes"):
                    lines.append(f"  Notes   : {p['other_notes']}")

        return "\n".join(lines)

    except Exception as e:
        return f"[Could not load loan catalog: {e}]"


# ══════════════════════════════════════════════════════════════════════════════
#  APPLICABILITY CHECK
# ══════════════════════════════════════════════════════════════════════════════
def _check_loan_applicability(extracted_text):
    """
    Fast pre-flight check: asks the LLM whether the extracted text looks like
    a loan application document.

    Returns (is_applicable: bool, reason: str)
      - is_applicable=True  → proceed with full analysis
      - is_applicable=False → reason explains exactly why it is not applicable
    """
    # Quick heuristic first — if text is very short it's probably garbage
    clean = extracted_text.strip()
    if len(clean) < 80:
        return False, "the extracted text is too short or empty to contain a loan application."

    prompt = f"""You are a document classifier for Banco San Vicente, a rural bank in the Philippines.

Your only job is to determine whether the document below is a loan application or loan-related document that can be meaningfully analyzed for credit assessment.

LOAN-RELATED documents include:
- Loan application forms (any type)
- Credit investigation / background investigation (CI/BI) forms
- Cashflow statements or income statements submitted with a loan
- Approval forms, credit approval sheets
- Financial statements (balance sheet, income statement) of a borrower
- Collateral appraisal reports
- Any form that is clearly part of a loan application package

NOT LOAN-RELATED documents include:
- General business documents unrelated to lending (e.g. sales reports, inventory lists)
- Medical records, school records, utility bills alone
- Random images, receipts, or personal notes with no financial/loan context
- Blank or near-blank documents
- Documents from a completely unrelated industry with no borrower/loan data

Respond ONLY in this exact JSON format (no explanation, no markdown):
{{"applicable": true, "reason": ""}}
or
{{"applicable": false, "reason": "brief specific reason why this document cannot be used for loan analysis"}}

--- DOCUMENT TEXT (first 2000 chars) ---
{clean[:2000]}
--- END ---
"""

    client = _get_groq()
    resp   = client.chat.completions.create(
        model       = "llama-3.3-70b-versatile",
        messages    = [{"role": "user", "content": prompt}],
        max_tokens  = 120,
        temperature = 0.0
    )
    raw = resp.choices[0].message.content.strip()

    # Strip markdown fences if model wraps it
    raw = re.sub(r'^```[a-z]*\n?', '', raw)
    raw = re.sub(r'\n?```$', '', raw)

    try:
        data = json.loads(raw)
        applicable = bool(data.get("applicable", True))
        reason     = str(data.get("reason", "")).strip()
        return applicable, reason
    except Exception:
        # If parsing fails, be lenient and allow the analysis to proceed
        return True, ""



def _run_credit_analysis(extracted_text):
    """
    Send extracted loan document text to LLM for credit assessment.
    - Gracefully handles missing / all-zero credit scoring sheet.
    - Injects the full BSV loan product catalog for product matching.
    - Produces a 10-section report: Section 1 leads with recommended
      products, upgrade path, and re-application guidance; followed by
      the full credit assessment (applicant summary → risk rating).
    Returns formatted analysis string.
    """
    # ── Detect if credit scoring sheet is empty / all-zero ────────────────
    scoring_available = True
    scoring_block     = ""
    in_scoring        = False
    for line in extracted_text.splitlines():
        if "CREDIT SCORING" in line.upper():
            in_scoring = True
        elif in_scoring and line.startswith("=== SHEET:"):
            break
        if in_scoring:
            scoring_block += line + "\n"

    if scoring_block:
        numbers  = re.findall(r'\b(\d+(?:\.\d+)?)\b', scoring_block)
        non_zero = [n for n in numbers if float(n) != 0.0]
        if not non_zero:
            scoring_available = False

    scoring_note = (
        "NOTE: The Credit Scoring sheet exists but contains all zeros — "
        "it has NOT been filled out by the account officer. "
        "Do NOT use it for scoring. Base your assessment solely on the "
        "financial data from the CI/BI and Approval Form sheets.\n\n"
        if not scoring_available else
        "The Credit Scoring sheet has been filled and may be used in assessment.\n\n"
    )

    # ── Build catalog context ─────────────────────────────────────────────
    loan_catalog_text = _build_loan_catalog_text()

    # ── Prompt ───────────────────────────────────────────────────────────
    prompt = f"""You are a senior credit analyst at Banco San Vicente, a rural bank in Camarines Norte, Philippines.
Analyze the following loan application data and produce a structured credit assessment report.

{scoring_note}
INSTRUCTIONS — produce ALL of the following numbered sections IN ORDER:

1. LOAN PRODUCT RECOMMENDATIONS
   Using the BANCO SAN VICENTE LOAN PRODUCT CATALOG provided below,
   match the applicant's profile (occupation, income, collateral, credit standing)
   to the most suitable BSV loan products.

   A) RECOMMENDED PRODUCTS (top 1–3 best matches)
      • Product name
      • Why it fits this applicant's profile
      • Estimated loanable amount in Philippine Pesos
      • Applicable interest rate and term
      • Documents still needed before submission

   B) UPGRADE PATH  (if the applicant does not fully qualify for their requested loan)
      • Specific steps they must take (build credit history, add collateral, etc.)
      • Which product to apply for now as a stepping stone
      • Realistic timeline to reach eligibility for the desired loan

   C) IF DECLINED ENTIRELY — guidance before re-applying
      • Minimum requirements currently missing
      • Concrete actions with approximate timeframes
      • Which BSV product to target first upon re-application

2. APPLICANT SUMMARY
   - Full name, age, occupation, civil status, address

3. FINANCIAL POSITION
   - Assets, liabilities, net worth
   - Debt-to-Asset Ratio  = Total Liabilities / Total Assets
   - Loan-to-Asset Ratio  = Loan Amount / Total Assets
   - Loan-to-Networth Ratio = Loan Amount / Net Worth

4. INCOME ANALYSIS
   - List all income sources with amounts
   - Monthly net income after expenses
   - If cashflow sheet is blank, note it clearly

5. EXISTING OBLIGATIONS
   - List all existing loans, monthly amortization burden, remaining balance

6. COLLATERAL ASSESSMENT
   - Real properties or assets offered, estimated values

7. CREDIT SCORING
   - Use rubric from document if available. If unavailable, state clearly and skip.

8. RISK FLAGS
   - Missing documents, data inconsistencies, concerns, or red flags

9. RECOMMENDATION
   - Choose one: APPROVE / CONDITIONALLY APPROVE / DECLINE
   - If CONDITIONALLY APPROVE: list the exact conditions required
   - Brief justification (3–5 sentences)

10. OVERALL RISK RATING
    - Low / Moderate / High with brief reason

Be specific with peso amounts. Use Philippine banking terminology.
Reference exact product names from the catalog.

--- LOAN APPLICATION DATA ---
{extracted_text}
--- END OF DATA ---

--- BANCO SAN VICENTE LOAN PRODUCT CATALOG ---
{loan_catalog_text}
--- END OF CATALOG ---
"""

    client = _get_groq()
    resp   = client.chat.completions.create(
        model       = "llama-3.3-70b-versatile",
        messages    = [{"role": "user", "content": prompt}],
        max_tokens  = 3000,
        temperature = 0.2
    )
    return resp.choices[0].message.content


# ══════════════════════════════════════════════════════════════════════════════
#  IMAGE PREPROCESSOR
# ══════════════════════════════════════════════════════════════════════════════
def _preprocess_image(file_path):
    from PIL import Image, ImageEnhance, ImageFilter
    import numpy as np

    img = Image.open(file_path).convert("RGB")
    w, h = img.size

    if w < 1800:
        scale = 1800 / w
        img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

    img = ImageEnhance.Sharpness(img).enhance(2.0)
    img = ImageEnhance.Contrast(img).enhance(1.5)
    img = ImageEnhance.Brightness(img).enhance(1.1)
    img = img.filter(ImageFilter.MedianFilter(size=3))

    try:
        import cv2
        arr  = np.array(img)
        gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
        _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
        coords = np.column_stack(np.where(binary > 0))
        if len(coords) > 100:
            angle = cv2.minAreaRect(coords.astype(np.float32))[-1]
            if angle < -45:
                angle = 90 + angle
            if abs(angle) > 0.5:
                (rh, rw) = gray.shape
                center   = (rw // 2, rh // 2)
                M        = cv2.getRotationMatrix2D(center, angle, 1.0)
                rotated  = cv2.warpAffine(arr, M, (rw, rh),
                                          flags=cv2.INTER_CUBIC,
                                          borderMode=cv2.BORDER_REPLICATE)
                img = Image.fromarray(rotated)
    except Exception:
        pass

    tmp_color = str(SCRIPT_DIR / "temp_preprocessed.jpg")
    tmp_gray  = str(SCRIPT_DIR / "temp_preprocessed_gray.jpg")
    img.save(tmp_color, "JPEG", quality=95)

    try:
        import cv2
        import numpy as np
        arr    = np.array(img)
        gray   = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
        binary = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, 31, 10
        )
        binary = cv2.medianBlur(binary, 3)
        cv2.imwrite(tmp_gray, binary)
    except Exception:
        tmp_gray = tmp_color

    return tmp_color, tmp_gray


# ══════════════════════════════════════════════════════════════════════════════
#  PARSERS
# ══════════════════════════════════════════════════════════════════════════════
def _parse_image(file_path, progress_cb=None):
    import numpy as np
    from PIL import Image

    if progress_cb: progress_cb(5,  "Stage 1: Preprocessing image…")
    tmp_color, tmp_gray = _preprocess_image(file_path)

    if progress_cb: progress_cb(15, "Stage 2: Running PaddleOCR…")
    ocr_text = ""
    try:
        arr = np.array(Image.open(tmp_gray).convert("RGB"))
        res = _get_paddle_ocr().ocr(arr, cls=True)
        if res and res[0]:
            for line in res[0]:
                if line[1][1] > 0.5:
                    ocr_text += line[1][0] + "\n"
    except Exception:
        ocr_text = ""

    if tmp_gray != tmp_color and os.path.exists(tmp_gray):
        os.remove(tmp_gray)

    if progress_cb: progress_cb(35, "Stage 3: Sending to Llama 4 Scout VLM…")

    ext = os.path.splitext(tmp_color)[1].lower()
    with open(tmp_color, "rb") as f:
        img_b64 = base64.b64encode(f.read()).decode()
    if os.path.exists(tmp_color):
        os.remove(tmp_color)

    mtype = ("image/jpeg" if ext in (".jpg", ".jpeg") else
             "image/png"  if ext == ".png"            else
             "image/bmp"  if ext == ".bmp"            else "image/tiff")

    ref = (f"\nOCR REFERENCE (PaddleOCR pre-read):\n{ocr_text}\n"
           if ocr_text.strip() else "")

    vlm_prompt = (
        "You are an expert document transcription assistant for a bank.\n"
        "This is a scanned loan-related form or document.\n\n"
        "INSTRUCTIONS:\n"
        "1. Transcribe ALL visible text exactly as it appears.\n"
        "2. Format each field as:  Field Label: value\n"
        "3. Empty fields → [EMPTY]\n"
        "4. Checked checkboxes → [CHECKED], unchecked → [UNCHECKED]\n"
        "5. If text is illegible or partially visible → [UNCLEAR]\n"
        "6. Preserve table structure: use  |  to separate columns.\n"
        "7. Handle multi-column layouts: read left column fully, then right.\n"
        "8. Read order: left→right, top→bottom.\n"
        "9. Include ALL numbers, amounts, dates, ID numbers exactly.\n"
        "10. Do NOT guess or fabricate values. If unsure, write [UNCLEAR].\n"
        + ref
    )

    client = _get_groq()
    resp   = client.chat.completions.create(
        model    = "meta-llama/llama-4-scout-17b-16e-instruct",
        messages = [{"role": "user", "content": [
            {"type": "image_url",
             "image_url": {"url": f"data:{mtype};base64,{img_b64}"}},
            {"type": "text", "text": vlm_prompt}
        ]}],
        max_tokens = 4000
    )

    if progress_cb: progress_cb(75, "Stage 3: VLM complete…")
    text = resp.choices[0].message.content

    if progress_cb: progress_cb(85, "Stage 4: Confidence check…")
    try:
        cr = client.chat.completions.create(
            model    = "llama-3.3-70b-versatile",
            messages = [{"role": "user", "content":
                f"Rate completeness 0-100. Reply ONLY a number.\n\n{text}"}],
            max_tokens=5, temperature=0
        )
        conf = int(cr.choices[0].message.content.strip())
    except Exception:
        conf = 50

    if progress_cb: progress_cb(95, f"Done — completeness: {conf}%")
    warn = (f"\n⚠ Low completeness ({conf}%). Complex layout?\n\n" if conf < 40 else "")
    return f"[Completeness: {conf}%]{warn}\n{text}"


def _parse_pdf(file_path, progress_cb=None):
    import pdfplumber
    text = ""
    af   = ""
    with pdfplumber.open(file_path) as pdf:
        total = len(pdf.pages)
        if progress_cb: progress_cb(5, f"PDF: {total} page(s)…")
        for p in pdf.pages:
            text += p.extract_text() or ""
        for p in pdf.pages:
            if p.annots:
                for a in p.annots:
                    if a.get("subtype") == "Widget" and a.get("V") and a.get("T"):
                        af += f"{a['T']}: {a['V']}\n"
        avg_chars = len(text) / max(total, 1)
        if avg_chars >= 50:
            if af:          return text + "\n=== FORM FIELDS ===\n" + af
            if text.strip(): return text

    text = ""
    dpi  = 300 if total <= 3 else 250 if total <= 10 else 200
    try:
        from pdf2image import convert_from_path
        kw = ({"poppler_path": POPPLER_PATH}
              if sys.platform == "win32" and Path(POPPLER_PATH).exists() else {})
        imgs = convert_from_path(file_path, dpi=dpi, fmt="jpeg",
                                  jpegopt={"quality": 95, "optimize": True},
                                  thread_count=min(4, total), **kw)
        for i, img in enumerate(imgs):
            if progress_cb:
                progress_cb(int(10 + (i / total) * 80), f"Page {i+1}/{total}…")
            tp = str(SCRIPT_DIR / f"tmp_pg_{i}.jpg")
            img.save(tp, "JPEG", quality=95)
            text += f"\n=== PAGE {i+1} ===\n{_parse_image(tp)}\n"
            os.remove(tp)
            del img
    except ImportError:
        text += "\n[pdf2image not installed]"
    except Exception as e:
        text += f"\n[PDF error: {e}]"
    return text


def _parse_excel(file_path, progress_cb=None):
    if progress_cb: progress_cb(10, "Reading Excel…")
    text = ""
    try:
        import openpyxl
        from datetime import datetime
        wb = openpyxl.load_workbook(file_path, data_only=True)

        def _get_merged_map(ws):
            merged = {}
            for rng in ws.merged_cells.ranges:
                label = f"[MERGED {rng.min_row},{rng.min_col}→{rng.max_row},{rng.max_col}]"
                for r in range(rng.min_row, rng.max_row + 1):
                    for c in range(rng.min_col, rng.max_col + 1):
                        merged[(r, c)] = label
            return merged

        def _format_cell(cell):
            val = cell.value
            if val is None: return ""
            nf = cell.number_format or ""
            if isinstance(val, datetime):
                return (val.strftime("%Y-%m-%d %H:%M")
                        if val.hour or val.minute
                        else val.strftime("%Y-%m-%d"))
            if isinstance(val, (int, float)):
                if "%" in nf:
                    return f"{val * 100:.2f}%"
                if any(c in nf for c in ("$","₱","€","£","¥")):
                    sym = next((c for c in ("$","₱","€","£","¥") if c in nf), "$")
                    return f"{sym}{val:,.2f}"
                if "," in nf and "." in nf:
                    return f"{val:,.2f}"
                if isinstance(val, float) and val == int(val):
                    return str(int(val))
            return str(val)

        for sn in wb.sheetnames:
            ws = wb[sn]
            if ws.sheet_state != "visible":
                text += f"\n=== {sn} [HIDDEN — SKIPPED] ===\n"
                continue
            text += f"\n=== SHEET: {sn} ===\n"
            merged_map = _get_merged_map(ws)
            if merged_map:
                text += f"  [Merged regions: {len(set(merged_map.values()))}]\n"
            for row in ws.iter_rows():
                cells_out = []
                for cell in row:
                    coord = (cell.row, cell.column)
                    if coord in merged_map:
                        if (cell.row == int(merged_map[coord].split(",")[0].split("[MERGED ")[1])
                                and cell.column == int(merged_map[coord].split("→")[0].split(",")[1])):
                            cells_out.append(_format_cell(cell))
                    else:
                        cells_out.append(_format_cell(cell))
                row_str = " | ".join(c for c in cells_out if c)
                if row_str.strip():
                    text += row_str + "\n"

        if progress_cb: progress_cb(95, "Done…")
        return text or "[Empty workbook]"

    except Exception as primary_err:
        if progress_cb: progress_cb(50, "openpyxl failed, trying pandas…")
        try:
            import pandas as pd
            text = f"[openpyxl error: {primary_err} — falling back to pandas]\n\n"
            xf   = pd.ExcelFile(file_path)
            for sn in xf.sheet_names:
                df = pd.read_excel(xf, sheet_name=sn, header=None)
                text += f"\n=== SHEET: {sn} ===\n"
                if df.empty:
                    text += "[Empty]\n"
                    continue
                for _, row in df.iterrows():
                    parts = [str(v).strip() for v in row
                             if str(v).strip() not in ("","nan")]
                    if parts: text += " | ".join(parts) + "\n"
            if progress_cb: progress_cb(95, "Done (pandas fallback)…")
            return text or "[Empty workbook]"
        except Exception as fallback_err:
            return (f"[Excel read failed]\n"
                    f"openpyxl: {primary_err}\npandas: {fallback_err}")


def _parse_docx(file_path, progress_cb=None):
    import docx as _d
    if progress_cb: progress_cb(10, "Reading Word…")
    doc   = _d.Document(file_path)
    text  = ""
    total = len(doc.paragraphs)
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip(): text += p.text + "\n"
        if progress_cb and i % 10 == 0:
            progress_cb(int(10 + (i / max(total,1)) * 70), f"Para {i+1}…")
    for t in doc.tables:
        for r in t.rows:
            rt = " | ".join(c.text.strip() for c in r.cells if c.text.strip())
            if rt: text += rt + "\n"
    if progress_cb: progress_cb(95, "Done…")
    return text or "[No text]"


def _parse_csv(file_path, progress_cb=None):
    import pandas as pd
    if progress_cb: progress_cb(10, "Reading CSV…")
    try:
        df = pd.read_csv(file_path)
    except Exception as e:
        return f"[CSV error: {e}]"
    if df.empty: return "[Empty CSV]"
    text = ""
    if len(df.columns) == 2:
        for _, r in df.iterrows():
            fld = str(r.iloc[0]).strip()
            val = str(r.iloc[1]).strip()
            if fld and fld.lower() != "nan":
                text += f"{fld}: {val if val.lower() != 'nan' else '[EMPTY]'}\n"
    else:
        text += "=== CSV ===\n" + "|".join(df.columns) + "\n\n"
        for i, r in df.iterrows():
            text += f"--- {i+1} ---\n"
            for c in df.columns:
                text += f"{c}: {str(r[c]).strip()}\n"
            text += "\n"
    if progress_cb: progress_cb(95, "Done…")
    return text or "[No data]"


def _parse_text(file_path, progress_cb=None):
    if progress_cb: progress_cb(20, "Reading…")
    for enc in ("utf-8","utf-8-sig","latin-1","cp1252"):
        try:
            t = Path(file_path).read_text(encoding=enc).strip()
            if progress_cb: progress_cb(90, "Done…")
            return t
        except (UnicodeDecodeError, LookupError):
            continue
    return "[Encoding error]"


def extract(file_path, progress_cb=None):
    ext = Path(file_path).suffix.lower()
    try:
        if ext in IMAGE_EXTS:           return _parse_image(file_path, progress_cb)
        elif ext == ".pdf":             return _parse_pdf(file_path, progress_cb)
        elif ext in (".docx",".doc"):   return _parse_docx(file_path, progress_cb)
        elif ext in (".xlsx",".xls"):   return _parse_excel(file_path, progress_cb)
        elif ext == ".csv":             return _parse_csv(file_path, progress_cb)
        else:                           return _parse_text(file_path, progress_cb)
    except EnvironmentError as e:
        return f"⚠ Config error:\n{e}"
    except ImportError as e:
        pkg = str(e).split("'")[1] if "'" in str(e) else str(e)
        return f"⚠ Missing: {pkg}\n\npip install {pkg}"
    except Exception as e:
        return f"⚠ Error:\n{type(e).__name__}: {e}"


# ══════════════════════════════════════════════════════════════════════════════
#  GRADIENT CANVAS
# ══════════════════════════════════════════════════════════════════════════════
class GradientCanvas(tk.Canvas):
    def __init__(self, parent, c1, c2, steps=60, **kw):
        super().__init__(parent, highlightthickness=0, bd=0, **kw)
        self.c1, self.c2, self.steps = c1, c2, steps
        self.bind("<Configure>", self._draw)

    def _draw(self, e=None):
        self.delete("g")
        w, h = self.winfo_width(), self.winfo_height()
        if w < 2 or h < 2: return
        sh = h / self.steps
        for i in range(self.steps):
            c = _hex_blend(self.c1, self.c2, i / self.steps)
            self.create_rectangle(0, int(i*sh), w, int((i+1)*sh)+1,
                                  fill=c, outline="", tags="g")
        self.lower("g")


# ══════════════════════════════════════════════════════════════════════════════
#  SPINNER
# ══════════════════════════════════════════════════════════════════════════════
class Spinner(tk.Canvas):
    def __init__(self, parent, size=96, bg=CARD_WHITE, **kw):
        super().__init__(parent, width=size, height=size,
                         bg=bg, highlightthickness=0, **kw)
        self.size  = size
        self.angle = 0
        self._job  = None

    def _draw(self):
        self.delete("all")
        cx = cy = self.size / 2
        r  = cx - 12
        self.create_oval(cx-r, cy-r, cx+r, cy+r, outline=BORDER_LIGHT, width=8)
        self.create_arc(cx-r, cy-r, cx+r, cy+r,
                        start=self.angle % 360, extent=300,
                        outline=NAVY, width=8, style="arc")
        self.create_arc(cx-r, cy-r, cx+r, cy+r,
                        start=(self.angle+300) % 360, extent=60,
                        outline=NAVY_GHOST, width=8, style="arc")
        rad = math.radians(self.angle % 360)
        dx  = cx + r * math.cos(rad)
        dy  = cy - r * math.sin(rad)
        self.create_oval(dx-7, dy-7, dx+7, dy+7,
                         fill=NAVY_MID, outline=NAVY_PALE, width=2)
        self.create_oval(dx-3, dy-3, dx+3, dy+3, fill=WHITE, outline="")

    def start(self): self._spin()
    def stop(self):
        if self._job:
            self.after_cancel(self._job)
            self._job = None
        self.delete("all")
    def _spin(self):
        self.angle += 5
        self._draw()
        self._job = self.after(14, self._spin)


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN APPLICATION
# ══════════════════════════════════════════════════════════════════════════════
class DocExtractorApp(ctk.CTk):

    def __init__(self):
        super().__init__()
        _register_fonts()
        global _FONT_FAMILY
        _FONT_FAMILY = _best_font()

        self.title("DocExtract Pro — Banco San Vicente")
        self.configure(fg_color=PANEL_LEFT)

        self.update_idletasks()
        x = (self.winfo_screenwidth()  - WIN_W) // 2
        y = (self.winfo_screenheight() - WIN_H) // 2
        self.geometry(f"{WIN_W}x{WIN_H}+{x}+{y}")
        self.resizable(False, False)

        self.overrideredirect(True)
        self._drag_x = self._drag_y = 0

        self._selected_file  = None
        self._logo_img       = None
        self._extracted_text = ""

        # ── Search state ──────────────────────────────────────────────────
        self._search_matches = []   # list of (start, end) index strings
        self._search_cursor  = -1   # which match is currently highlighted
        self._last_query     = ""   # last searched term

        self._build_ui()

        self.after(100, self._fix_windows_taskbar)
        self.after(100, self._force_focus)

    # ── Windows taskbar fix ────────────────────────────────────────────────
    def _fix_windows_taskbar(self):
        if sys.platform == "win32":
            GWL_EXSTYLE      = -20
            WS_EX_APPWINDOW  = 0x00040000
            WS_EX_TOOLWINDOW = 0x00000080
            hwnd  = ctypes.windll.user32.GetParent(self.winfo_id())
            style = ctypes.windll.user32.GetWindowLongW(hwnd, GWL_EXSTYLE)
            style = (style & ~WS_EX_TOOLWINDOW) | WS_EX_APPWINDOW
            ctypes.windll.user32.SetWindowLongW(hwnd, GWL_EXSTYLE, style)
            self.withdraw()
            self.after(10, self.deiconify)

    def _force_focus(self):
        self.lift()
        self.attributes("-topmost", True)
        self.after(300, lambda: self.attributes("-topmost", False))
        self.focus_force()

    def _do_minimize(self):
        if sys.platform == "win32":
            hwnd = ctypes.windll.user32.GetParent(self.winfo_id())
            ctypes.windll.user32.ShowWindow(hwnd, 6)  # SW_MINIMIZE
        else:
            self.iconify()

    def _on_restore(self, e):
        self.unbind("<Map>")
        self._force_focus()

    # ── UI BUILD ───────────────────────────────────────────────────────────
    def _build_ui(self):
        self._build_topbar()
        body = tk.Frame(self, bg=PANEL_LEFT)
        body.pack(fill="both", expand=True)

        left = tk.Frame(body, bg=PANEL_LEFT, width=310)
        left.pack(side="left", fill="y")
        left.pack_propagate(False)

        tk.Frame(body, bg=BORDER_LIGHT, width=1).pack(side="left", fill="y")

        right = tk.Frame(body, bg=CARD_WHITE)
        right.pack(side="left", fill="both", expand=True)

        self._build_left(left)
        self._build_right(right)

    def _build_topbar(self):
        bar = GradientCanvas(self, NAVY_DEEP, NAVY_MID, steps=50, height=52)
        bar.pack(fill="x")
        bar.bind("<ButtonPress-1>", self._drag_start)
        bar.bind("<B1-Motion>",     self._drag_move)

        close_btn = tk.Label(bar, text="  ✕  ", font=F(11,"bold"),
                              fg=WHITE, bg="#C0392B", cursor="hand2",
                              width=3, anchor="center")
        close_btn.bind("<Enter>",    lambda e: close_btn.config(bg="#E74C3C"))
        close_btn.bind("<Leave>",    lambda e: close_btn.config(bg="#C0392B"))
        close_btn.bind("<Button-1>", lambda e: self.destroy())
        close_btn.place(in_=bar, relx=1.0, rely=0.0, anchor="ne",
                        relheight=1.0, width=46)

        min_btn = tk.Label(bar, text="  —  ", font=F(11,"bold"),
                            fg=WHITE, bg=NAVY_MID, cursor="hand2",
                            width=3, anchor="center")
        min_btn.bind("<Enter>",    lambda e: min_btn.config(bg=NAVY_LIGHT))
        min_btn.bind("<Leave>",    lambda e: min_btn.config(bg=NAVY_MID))
        min_btn.bind("<Button-1>", lambda e: self._do_minimize())
        min_btn.place(in_=bar, relx=1.0, rely=0.0, anchor="ne",
                      relheight=1.0, width=46, x=-46)

        acc = tk.Canvas(self, height=4, highlightthickness=0)
        acc.pack(fill="x")
        acc.bind("<Configure>",
                 lambda e, c=acc: self._hbar(c, NAVY, LIME_BRIGHT, 70))

    def _drag_start(self, e):
        self._drag_x = e.x_root - self.winfo_x()
        self._drag_y = e.y_root - self.winfo_y()

    def _drag_move(self, e):
        self.geometry(f"+{e.x_root - self._drag_x}+{e.y_root - self._drag_y}")

    def _build_left(self, p):
        wrap = tk.Frame(p, bg=PANEL_LEFT)
        wrap.pack(fill="both", expand=True, padx=22)

        logo_loaded = False
        if LOGO_PATH.exists():
            try:
                from PIL import Image
                img    = Image.open(LOGO_PATH).convert("RGBA")
                bg_pil = Image.new("RGBA", img.size,
                                   tuple(int(PANEL_LEFT[i:i+2],16)
                                         for i in (1,3,5)) + (255,))
                bg_pil.paste(img, mask=img.split()[3])
                bg_pil = bg_pil.convert("RGB")
                bg_pil.thumbnail((200, 48), Image.LANCZOS)
                self._logo_img = ctk.CTkImage(
                    light_image=bg_pil, dark_image=bg_pil,
                    size=(bg_pil.width, bg_pil.height)
                )
                lf = ctk.CTkFrame(wrap, fg_color="transparent")
                lf.pack(pady=(20,6))
                ctk.CTkLabel(lf, image=self._logo_img, text="",
                              fg_color="transparent").pack()
                logo_loaded = True
            except Exception:
                pass

        if not logo_loaded:
            mark_frame = tk.Frame(wrap, bg=PANEL_LEFT)
            mark_frame.pack(pady=(20,6))
            mark = tk.Canvas(mark_frame, width=34, height=38,
                              bg=PANEL_LEFT, highlightthickness=0)
            mark.pack(side="left", padx=(0,10))
            mark.create_rectangle(0,12,18,38, fill=NAVY_MID,    outline="")
            mark.create_rectangle(16, 0,34,26, fill=LIME_BRIGHT, outline="")
            tk.Label(mark_frame, text="Banco San Vicente",
                      font=F(11,"bold"), fg=NAVY_DEEP, bg=PANEL_LEFT).pack(side="left")

        self._div(wrap)

        title_row = tk.Frame(wrap, bg=PANEL_LEFT)
        title_row.pack(fill="x")
        acc = tk.Canvas(title_row, width=5, height=58,
                         bg=PANEL_LEFT, highlightthickness=0)
        acc.pack(side="left")
        acc.bind("<Configure>", lambda e, c=acc: self._vbar(c, LIME_BRIGHT, LIME_DARK))
        txt = tk.Frame(title_row, bg=PANEL_LEFT)
        txt.pack(side="left", padx=12)
        tk.Label(txt, text="MaYuKen — OCR",
                  font=F(17,"bold"), fg=NAVY_DEEP, bg=PANEL_LEFT,
                  wraplength=240, anchor="w", justify="left").pack(anchor="w")
        tk.Label(txt, text="Document & Image Extraction",
                  font=F(8), fg=TXT_SOFT, bg=PANEL_LEFT).pack(anchor="w", pady=(2,0))

        self._div(wrap)
        self._sec(wrap, "UPLOAD FILE")

        drop = tk.Frame(wrap, bg=NAVY_MIST,
                         highlightbackground=BORDER_MID, highlightthickness=1, height=110)
        drop.pack(fill="x")
        drop.pack_propagate(False)
        stripe = tk.Canvas(drop, height=4, bg=NAVY_MIST, highlightthickness=0)
        stripe.place(x=0, y=0, relwidth=1)
        stripe.bind("<Configure>",
                    lambda e, c=stripe: self._hbar(c, NAVY_LIGHT, LIME))
        inner = tk.Frame(drop, bg=NAVY_MIST)
        inner.place(relx=0.5, rely=0.5, anchor="center")
        self._icon_lbl = tk.Label(inner, text="📁",
                                   font=("Segoe UI Emoji",26),
                                   fg=NAVY_MID, bg=NAVY_MIST)
        self._icon_lbl.pack()
        self._filename_lbl = tk.Label(inner, text="No file selected",
                                       font=F(8), fg=TXT_MUTED, bg=NAVY_MIST,
                                       wraplength=230, justify="center")
        self._filename_lbl.pack(pady=(4,0))

        tk.Frame(wrap, bg=PANEL_LEFT, height=10).pack()

        self._browse_btn = ctk.CTkButton(
            wrap, text="  Browse File", command=self._browse,
            height=44, corner_radius=22, fg_color=NAVY, hover_color=NAVY_LIGHT,
            text_color=WHITE, font=ctk.CTkFont(_FONT_FAMILY, 11, weight="bold"),
            border_width=0
        )
        self._browse_btn.pack(fill="x", pady=(0,7))

        self._ext_btn = ctk.CTkButton(
            wrap, text="  Extract Text", command=self._start_extraction,
            height=44, corner_radius=22, fg_color=LIME, hover_color=LIME_BRIGHT,
            text_color=TXT_ON_LIME, font=ctk.CTkFont(_FONT_FAMILY, 11, weight="bold"),
            state="disabled", border_width=0
        )
        self._ext_btn.pack(fill="x", pady=(0,7))

        self._analyze_btn = ctk.CTkButton(
            wrap, text="  Analyze Loan",
            command=self._start_analysis,
            height=44, corner_radius=22,
            fg_color=NAVY_DEEP, hover_color=NAVY_MID,
            text_color=WHITE,
            font=ctk.CTkFont(_FONT_FAMILY, 11, weight="bold"),
            state="disabled", border_width=0
        )
        self._analyze_btn.pack(fill="x")

        self._div(wrap)
        self._sec(wrap, "OCR ENGINE")
        for icon, label, col in [
            ("🤖", "PaddleOCR  (Stage 1)",          NAVY),
            ("🦙", "Llama 4 Scout VLM  (Stage 2)",  LIME_DARK),
            ("✅", "Confidence Check  (Stage 3)",    ACCENT_SUCCESS),
            ("🏦", "Credit Analysis  (LLaMA 3.3)",  NAVY_DEEP),
        ]:
            row = tk.Frame(wrap, bg=PANEL_LEFT)
            row.pack(fill="x", pady=2)
            tk.Label(row, text=icon, font=("Segoe UI Emoji",10),
                      fg=col, bg=PANEL_LEFT).pack(side="left")
            tk.Label(row, text=label, font=F(8), fg=TXT_SOFT,
                      bg=PANEL_LEFT).pack(side="left", padx=(9,0))

        self._div(wrap)
        self._sec(wrap, "SUPPORTED FORMATS")
        for icon, label in [
            ("📄","PDF  (.pdf)"),("📝","Word  (.docx)"),
            ("📊","Excel  (.xlsx)"),("📃","Text / CSV / MD"),
            ("🖼","Images  (.png .jpg .bmp .tiff .webp)"),
        ]:
            row = tk.Frame(wrap, bg=PANEL_LEFT)
            row.pack(fill="x", pady=2)
            tk.Label(row, text=icon, font=("Segoe UI Emoji",10),
                      fg=NAVY_PALE, bg=PANEL_LEFT).pack(side="left")
            tk.Label(row, text=label, font=F(8), fg=TXT_MUTED,
                      bg=PANEL_LEFT).pack(side="left", padx=(9,0))

    def _build_right(self, p):
        hdr = tk.Frame(p, bg=CARD_WHITE)
        hdr.pack(fill="x", padx=32, pady=(26,0))

        tk.Label(hdr, text="Extracted Content",
                  font=F(20,"bold"), fg=NAVY_DEEP, bg=CARD_WHITE).pack(side="left")

        badge_bg = NAVY_MIST
        badge    = tk.Frame(hdr, bg=badge_bg,
                             highlightbackground=BORDER_MID, highlightthickness=1)
        badge.pack(side="left", padx=(16,0), pady=2)
        self._status_lbl = tk.Label(badge, text="●  Ready",
                                     font=F(8,"bold"), fg=LIME_DARK,
                                     bg=badge_bg, padx=14, pady=5)
        self._status_lbl.pack()

        self._copy_btn = ctk.CTkButton(
            hdr, text="⎘  Copy All", command=self._copy,
            width=120, height=34, corner_radius=17,
            fg_color=NAVY_MIST, hover_color=NAVY_GHOST, text_color=NAVY,
            font=ctk.CTkFont(_FONT_FAMILY, 9, weight="bold"),
            border_width=1, border_color=BORDER_MID
        )
        self._copy_btn.pack(side="right")

        # ── Tab strip ─────────────────────────────────────────────────────
        tab_row = tk.Frame(p, bg=CARD_WHITE)
        tab_row.pack(fill="x", padx=32, pady=(10,0))

        self._active_tab = tk.StringVar(value="extract")

        def _tab_style(btn, active):
            if active:
                btn.configure(fg_color=NAVY,      text_color=WHITE,   hover_color=NAVY_LIGHT)
            else:
                btn.configure(fg_color=NAVY_MIST, text_color=TXT_SOFT, hover_color=NAVY_GHOST)

        self._tab_extract_btn = ctk.CTkButton(
            tab_row, text="📄  Extracted Text", width=160, height=30,
            corner_radius=15, font=ctk.CTkFont(_FONT_FAMILY, 9, weight="bold"),
            border_width=0, command=lambda: self._switch_tab("extract")
        )
        self._tab_extract_btn.pack(side="left", padx=(0,6))

        self._tab_analysis_btn = ctk.CTkButton(
            tab_row, text="🏦  Loan Analysis", width=160, height=30,
            corner_radius=15, font=ctk.CTkFont(_FONT_FAMILY, 9, weight="bold"),
            border_width=0, command=lambda: self._switch_tab("analysis")
        )
        self._tab_analysis_btn.pack(side="left")

        _tab_style(self._tab_extract_btn,  True)
        _tab_style(self._tab_analysis_btn, False)
        self._tab_style_fn = _tab_style

        tk.Label(
            p,
            text=("Hybrid Pipeline:  PaddleOCR  →  Llama 4 Scout VLM  "
                  "→  Confidence Scoring  →  Credit & Eligibility Analysis"),
            font=F(8), fg=TXT_MUTED, bg=CARD_WHITE
        ).pack(anchor="w", padx=32, pady=(6,4))

        # ── Search bar ────────────────────────────────────────────────────
        search_row = tk.Frame(p, bg=CARD_WHITE)
        search_row.pack(fill="x", padx=32, pady=(0,8))

        # search icon + entry
        search_wrap = tk.Frame(search_row, bg=NAVY_MIST,
                                highlightbackground=BORDER_MID, highlightthickness=1)
        search_wrap.pack(side="left", fill="x", expand=True)

        tk.Label(search_wrap, text="🔍", font=("Segoe UI Emoji", 10),
                  bg=NAVY_MIST, fg=NAVY_MID).pack(side="left", padx=(8,2))

        self._search_var = tk.StringVar()
        self._search_entry = tk.Entry(
            search_wrap, textvariable=self._search_var,
            font=F(9), fg=TXT_NAVY, bg=NAVY_MIST,
            relief="flat", bd=0, insertbackground=NAVY,
            width=28
        )
        self._search_entry.pack(side="left", fill="x", expand=True, pady=6)
        self._search_entry.bind("<Return>",   lambda e: self._do_search())
        self._search_entry.bind("<KP_Enter>", lambda e: self._do_search())
        self._search_entry.bind("<Escape>",   lambda e: self._clear_search())
        self._search_var.trace_add("write",   lambda *a: self._do_search())

        # match counter label
        self._match_lbl = tk.Label(search_wrap, text="",
                                    font=F(8), fg=TXT_SOFT, bg=NAVY_MIST,
                                    padx=8)
        self._match_lbl.pack(side="left")

        # nav buttons
        nav_frame = tk.Frame(search_row, bg=CARD_WHITE)
        nav_frame.pack(side="left", padx=(6,0))

        self._prev_btn = ctk.CTkButton(
            nav_frame, text="▲", width=30, height=30, corner_radius=15,
            fg_color=NAVY_MIST, hover_color=NAVY_GHOST, text_color=NAVY,
            font=ctk.CTkFont(_FONT_FAMILY, 9, weight="bold"),
            border_width=1, border_color=BORDER_MID,
            command=self._search_prev
        )
        self._prev_btn.pack(side="left", padx=(0,4))

        self._next_btn = ctk.CTkButton(
            nav_frame, text="▼", width=30, height=30, corner_radius=15,
            fg_color=NAVY_MIST, hover_color=NAVY_GHOST, text_color=NAVY,
            font=ctk.CTkFont(_FONT_FAMILY, 9, weight="bold"),
            border_width=1, border_color=BORDER_MID,
            command=self._search_next
        )
        self._next_btn.pack(side="left", padx=(0,4))

        clr_btn = ctk.CTkButton(
            nav_frame, text="✕", width=30, height=30, corner_radius=15,
            fg_color=NAVY_MIST, hover_color="#FFD6D6", text_color=ACCENT_RED,
            font=ctk.CTkFont(_FONT_FAMILY, 9, weight="bold"),
            border_width=1, border_color=BORDER_MID,
            command=self._clear_search
        )
        clr_btn.pack(side="left")

        # ── Card ──────────────────────────────────────────────────────────
        card_outer = tk.Frame(p, bg=BORDER_LIGHT, padx=1, pady=1)
        card_outer.pack(fill="both", expand=True, padx=28, pady=(0,28))

        card = tk.Frame(card_outer, bg=CARD_WHITE)
        card.pack(fill="both", expand=True)

        top_acc = tk.Canvas(card, height=5, bg=CARD_WHITE, highlightthickness=0)
        top_acc.pack(fill="x")
        top_acc.bind("<Configure>",
                     lambda e, c=top_acc: self._hbar(c, NAVY, LIME_BRIGHT, 80))

        # ── Loader frame ──────────────────────────────────────────────────
        self._loader_frame = tk.Frame(card, bg=CARD_WHITE)
        tk.Frame(self._loader_frame, bg=CARD_WHITE).pack(expand=True, fill="both")
        center = tk.Frame(self._loader_frame, bg=CARD_WHITE)
        center.pack()
        self._spinner = Spinner(center, size=96, bg=CARD_WHITE)
        self._spinner.pack(pady=(0,20))
        tk.Label(center, text="Processing…",
                  font=F(15,"bold"), fg=NAVY_DEEP, bg=CARD_WHITE).pack()
        self._stage_lbl = tk.Label(center, text="Initialising…",
                                    font=F(9), fg=TXT_SOFT, bg=CARD_WHITE)
        self._stage_lbl.pack(pady=(6,2))
        self._pct_lbl = tk.Label(center, text="0%",
                                  font=F(13,"bold"), fg=LIME_DARK, bg=CARD_WHITE)
        self._pct_lbl.pack(pady=(0,16))
        self._prog_bar = ctk.CTkProgressBar(
            center, width=300, height=9, corner_radius=5,
            fg_color=NAVY_GHOST, progress_color=LIME, border_width=0
        )
        self._prog_bar.set(0)
        self._prog_bar.pack()
        tk.Frame(self._loader_frame, bg=CARD_WHITE).pack(expand=True, fill="both")

        # ── Extracted text panel ──────────────────────────────────────────
        self._txt_frame = tk.Frame(card, bg=CARD_WHITE)
        self._txt_frame.pack(fill="both", expand=True)
        sb = tk.Scrollbar(self._txt_frame, relief="flat", troughcolor=OFF_WHITE,
                          bg=BORDER_LIGHT, width=10, bd=0)
        sb.pack(side="right", fill="y")
        self._textbox = tk.Text(
            self._txt_frame, wrap="word", font=FMONO(11),
            fg=TXT_NAVY, bg=CARD_WHITE, relief="flat", bd=0,
            padx=30, pady=22, spacing1=4, spacing2=2, spacing3=4,
            insertbackground=LIME_DARK, yscrollcommand=sb.set,
            state="disabled", cursor="arrow",
            selectbackground=NAVY_GHOST, selectforeground=TXT_NAVY
        )
        self._textbox.pack(side="left", fill="both", expand=True)
        sb.config(command=self._textbox.yview)
        self._textbox.tag_configure("search_match",  background=LIME_PALE,   foreground=TXT_NAVY)
        self._textbox.tag_configure("search_current", background=LIME_BRIGHT, foreground=NAVY_DEEP)

        # ── Analysis panel ────────────────────────────────────────────────
        self._analysis_frame = tk.Frame(card, bg=CARD_WHITE)
        sb2 = tk.Scrollbar(self._analysis_frame, relief="flat", troughcolor=OFF_WHITE,
                           bg=BORDER_LIGHT, width=10, bd=0)
        sb2.pack(side="right", fill="y")
        self._analysis_box = tk.Text(
            self._analysis_frame, wrap="word", font=F(12),
            fg=TXT_NAVY, bg=CARD_WHITE, relief="flat", bd=0,
            padx=36, pady=26, spacing1=6, spacing2=3, spacing3=8,
            insertbackground=LIME_DARK, yscrollcommand=sb2.set,
            state="disabled", cursor="arrow",
            selectbackground=NAVY_GHOST, selectforeground=TXT_NAVY
        )
        self._analysis_box.pack(side="left", fill="both", expand=True)
        sb2.config(command=self._analysis_box.yview)
        self._configure_analysis_tags(self._analysis_box)

        self._put_placeholder()
        self._current_tab = "extract"

    # ── TAG CONFIGURATION — ANALYSIS ─────────────────────────────────────
    def _configure_analysis_tags(self, box):
        sz = 12
        ff = _FONT_FAMILY
        box.tag_configure("search_match",    background=LIME_PALE,   foreground=TXT_NAVY)
        box.tag_configure("search_current",  background=LIME_BRIGHT, foreground=NAVY_DEEP)
        # Section headers  e.g. "1. LOAN PRODUCT RECOMMENDATIONS"
        box.tag_configure("sec_header",      font=(ff, sz + 2, "bold"), foreground=NAVY_DEEP,
                           spacing1=18, spacing3=6, lmargin1=0, lmargin2=0)
        # Sub-section labels  e.g. "A) RECOMMENDED PRODUCTS"
        box.tag_configure("sub_header",      font=(ff, sz, "bold"),     foreground=NAVY_MID,
                           spacing1=10, spacing3=4, lmargin1=10, lmargin2=10)
        # Bullet items  "•  ..."
        box.tag_configure("bullet",          font=(ff, sz),             foreground=TXT_NAVY,
                           lmargin1=24, lmargin2=36, spacing1=3, spacing3=3)
        # Normal body text
        box.tag_configure("body",            font=(ff, sz),             foreground=TXT_NAVY,
                           lmargin1=10, lmargin2=10, spacing1=3, spacing3=3)
        # Approve / Decline / Conditionally Approve verdict
        box.tag_configure("verdict_approve", font=(ff, sz, "bold"),     foreground=ACCENT_SUCCESS)
        box.tag_configure("verdict_cond",    font=(ff, sz, "bold"),     foreground=ACCENT_GOLD)
        box.tag_configure("verdict_decline", font=(ff, sz, "bold"),     foreground=ACCENT_RED)
        # ✅ ⚠ ❌ symbols in body text
        box.tag_configure("sym_ok",          foreground=ACCENT_SUCCESS, font=(ff, sz))
        box.tag_configure("sym_warn",        foreground=ACCENT_GOLD,    font=(ff, sz))
        box.tag_configure("sym_bad",         foreground=ACCENT_RED,     font=(ff, sz))
        # Peso amounts ₱
        box.tag_configure("peso",            foreground=LIME_DARK,      font=(ff, sz, "bold"))
        # Risk rating
        box.tag_configure("risk_low",        foreground=ACCENT_SUCCESS, font=(ff, sz, "bold"))
        box.tag_configure("risk_mod",        foreground=ACCENT_GOLD,    font=(ff, sz, "bold"))
        box.tag_configure("risk_high",       foreground=ACCENT_RED,     font=(ff, sz, "bold"))
        # Divider lines
        box.tag_configure("divider",         foreground=BORDER_MID,     font=(ff, sz - 2))

    # ── WRITE — EXTRACTED TEXT (plain mono) ──────────────────────────────
    def _write(self, txt, color=TXT_NAVY):
        box = self._textbox
        box.config(state="normal", fg=color)
        box.delete("1.0", "end")
        box.insert("end", txt)
        box.config(state="disabled")
        if self._search_var.get().strip() and self._current_tab == "extract":
            self._do_search()

    # ── RICH WRITE — ANALYSIS ─────────────────────────────────────────────
    def _write_analysis(self, txt, color=TXT_NAVY):
        box = self._analysis_box
        box.config(state="normal")
        box.delete("1.0", "end")

        # If plain placeholder/error text, write simply
        if color != TXT_NAVY:
            box.insert("end", txt)
            box.config(state="disabled")
            return

        for line in txt.splitlines(keepends=True):
            s = line.rstrip("\n")

            # Numbered section headers  "1. LOAN PRODUCT..." / "10. OVERALL..."
            if re.match(r'^\d{1,2}\.\s+[A-Z]', s):
                box.insert("end", "\n" + s + "\n", "sec_header")

            # Sub-section labels  "A) ..." / "B) ..." / "C) ..."
            elif re.match(r'^[A-C]\)\s', s):
                box.insert("end", s + "\n", "sub_header")

            # Verdict lines
            elif re.search(r'\bAPPROVE\b', s) and not re.search(r'CONDITIONALLY', s):
                self._insert_with_peso(box, s + "\n", "verdict_approve")
            elif re.search(r'CONDITIONALLY APPROVE', s):
                self._insert_with_peso(box, s + "\n", "verdict_cond")
            elif re.search(r'\bDECLINE\b', s):
                self._insert_with_peso(box, s + "\n", "verdict_decline")

            # Risk rating
            elif re.search(r'\b(Low Risk|LOW RISK|Low)\b', s) and "risk" in s.lower():
                self._insert_with_peso(box, s + "\n", "risk_low")
            elif re.search(r'\b(Moderate Risk|MODERATE|Moderate)\b', s) and "risk" in s.lower():
                self._insert_with_peso(box, s + "\n", "risk_mod")
            elif re.search(r'\b(High Risk|HIGH RISK|High)\b', s) and "risk" in s.lower():
                self._insert_with_peso(box, s + "\n", "risk_high")

            # Divider lines
            elif set(s.strip()).issubset({"─","—","-","=","_","*"}) and len(s.strip()) > 4:
                box.insert("end", s + "\n", "divider")

            # Bullet items
            elif s.strip().startswith(("•", "-", "*", "–")) or re.match(r'^\s+[•\-\*]', s):
                self._insert_with_peso(box, s + "\n", "bullet")

            # Lines with ✅ ⚠ ❌
            elif "✅" in s:
                self._insert_sym_line(box, s + "\n", "sym_ok")
            elif "⚠" in s:
                self._insert_sym_line(box, s + "\n", "sym_warn")
            elif "❌" in s:
                self._insert_sym_line(box, s + "\n", "sym_bad")

            else:
                self._insert_with_peso(box, s + "\n", "body")

        box.config(state="disabled")
        if self._search_var.get().strip() and self._current_tab == "analysis":
            self._do_search()

    def _insert_with_peso(self, box, text, base_tag):
        """Insert text with base_tag, but colour ₱ amounts in peso tag."""
        parts = re.split(r'(₱[\d,]+(?:\.\d+)?)', text)
        for part in parts:
            if re.match(r'₱[\d,]+', part):
                box.insert("end", part, "peso")
            else:
                box.insert("end", part, base_tag)

    def _insert_sym_line(self, box, text, sym_tag):
        """Insert a line colouring the leading symbol differently from body."""
        # colour the first symbol, rest as body with peso highlights
        m = re.match(r'^(\s*[✅⚠❌]\s*)', text)
        if m:
            box.insert("end", m.group(1), sym_tag)
            self._insert_with_peso(box, text[m.end():], "body")
        else:
            self._insert_with_peso(box, text, "body")

    def _switch_tab(self, tab):
        self._current_tab = tab
        self._tab_style_fn(self._tab_extract_btn,  tab == "extract")
        self._tab_style_fn(self._tab_analysis_btn, tab == "analysis")
        self._loader_frame.pack_forget()
        self._txt_frame.pack_forget()
        self._analysis_frame.pack_forget()
        if tab == "extract":
            self._txt_frame.pack(fill="both", expand=True)
        else:
            self._analysis_frame.pack(fill="both", expand=True)
        # Re-apply search highlights on the newly visible tab
        if self._search_var.get().strip():
            self._do_search()

    def _hbar(self, canvas, c1, c2, steps=40):
        canvas.delete("all")
        w = canvas.winfo_width()
        h = canvas.winfo_height()
        if w < 2: return
        for i in range(steps):
            c = _hex_blend(c1, c2, i / steps)
            canvas.create_rectangle(int(w*i/steps), 0, int(w*(i+1)/steps)+1, h,
                                    fill=c, outline="")

    def _vbar(self, canvas, c1, c2, steps=20):
        canvas.delete("all")
        h = canvas.winfo_height()
        if h < 2: return
        for i in range(steps):
            c = _hex_blend(c1, c2, i / steps)
            canvas.create_rectangle(0, int(h*i/steps), 5, int(h*(i+1)/steps)+1,
                                    fill=c, outline="")

    def _div(self, parent):
        tk.Frame(parent, bg=BORDER_LIGHT, height=1).pack(fill="x", pady=(18,14))

    def _sec(self, parent, text):
        tk.Label(parent, text=text, font=F(7,"bold"),
                  fg=NAVY_PALE, bg=PANEL_LEFT).pack(anchor="w", pady=(0,7))

    def _put_placeholder(self):
        self._write(
            "Results will appear here after extraction.\n\n"
            "←  Choose a file from the left panel, then click  Extract Text  to begin.\n\n"
            "Engine:  PaddleOCR  →  Llama 4 Scout VLM  →  Confidence Check",
            color=TXT_MUTED
        )
        self._write_analysis(
            "Loan analysis will appear here.\n\n"
            "←  First extract a loan document, then click  Analyze Loan.\n\n"
            "The AI will assess creditworthiness, flag risks, give a recommendation,\n"
            "and match the applicant against all BSV loan products — including\n"
            "eligible amounts, upgrade paths, and re-application guidance.\n\n"
            "Note: Works even if the Credit Scoring sheet is blank.",
            color=TXT_MUTED
        )

    def _show_loader(self, show, stage_text="Processing…"):
        if show:
            self._txt_frame.pack_forget()
            self._analysis_frame.pack_forget()
            self._loader_frame.pack(fill="both", expand=True)
            self._stage_lbl.config(text=stage_text)
            self._spinner.start()
            self._status_lbl.config(text="●  Processing…", fg=ACCENT_GOLD)
        else:
            self._spinner.stop()
            self._loader_frame.pack_forget()
            if self._current_tab == "extract":
                self._txt_frame.pack(fill="both", expand=True)
            else:
                self._analysis_frame.pack(fill="both", expand=True)
            self._status_lbl.config(text="●  Ready", fg=LIME_DARK)

    def _set_progress(self, pct, stage=""):
        self._pct_lbl.config(text=f"{pct}%")
        if stage: self._stage_lbl.config(text=stage)
        self._prog_bar.set(pct / 100)

    def _file_icon_for(self, name):
        ext = Path(name).suffix.lower()
        if ext == ".pdf":              return "📄"
        if ext in (".docx",".doc"):    return "📝"
        if ext in (".xlsx",".xls"):    return "📊"
        if ext in IMAGE_EXTS:          return "🖼"
        return "📃"

    def _browse(self):
        path = filedialog.askopenfilename(
            title="Select a document or image",
            filetypes=[
                ("All supported","*.pdf *.txt *.docx *.xlsx *.xls *.csv *.md "
                 "*.png *.jpg *.jpeg *.bmp *.tiff *.tif *.webp *.gif"),
                ("PDF","*.pdf"),("Word","*.docx *.doc"),
                ("Excel","*.xlsx *.xls"),("Text/CSV","*.txt *.csv *.md"),
                ("Images","*.png *.jpg *.jpeg *.bmp *.tiff *.tif *.webp *.gif"),
                ("All files","*.*"),
            ]
        )
        if path:
            self._selected_file = path
            name  = path.replace("\\","/").split("/")[-1]
            short = name if len(name) <= 32 else name[:29] + "…"
            self._filename_lbl.config(text=short, fg=TXT_NAVY_MID)
            self._icon_lbl.config(text=self._file_icon_for(name), fg=LIME_DARK)
            self._ext_btn.configure(state="normal")

    def _start_extraction(self):
        if not self._selected_file: return
        self._ext_btn.configure(state="disabled")
        self._analyze_btn.configure(state="disabled")
        self._switch_tab("extract")
        self._show_loader(True, "Extracting document…")
        self._set_progress(0, "Starting…")

        def worker():
            def cb(pct, stage=""): self.after(0, self._set_progress, pct, stage)
            result = extract(self._selected_file, cb)
            self.after(0, self._finish_extraction, result)

        threading.Thread(target=worker, daemon=True).start()

    def _finish_extraction(self, result):
        self._extracted_text = result
        self._show_loader(False)
        name = Path(self._selected_file).name
        ext  = Path(self._selected_file).suffix.upper().lstrip(".")
        hdr  = (
            f"File   : {name}\n"
            f"Type   : {ext}\n"
            f"Chars  : {len(result):,}\n"
            f"Lines  : {result.count(chr(10)):,}\n"
            + "─" * 58 + "\n\n"
        )
        self._write(hdr + result, TXT_NAVY)
        self._status_lbl.config(text="●  Complete", fg=LIME_DARK)
        self._ext_btn.configure(state="normal")
        self._analyze_btn.configure(state="normal")

    # ── ANALYSIS ──────────────────────────────────────────────────────────
    def _start_analysis(self):
        if not self._extracted_text.strip():
            self._write_analysis(
                "⚠ No extracted text found.\n\nPlease extract a document first.",
                color=ACCENT_RED
            )
            self._switch_tab("analysis")
            return

        self._analyze_btn.configure(state="disabled")
        self._switch_tab("analysis")
        self._show_loader(True, "Checking document…")
        self._set_progress(0, "Verifying document type…")

        def worker():
            # ── Step 1: applicability check ───────────────────────────────
            try:
                self.after(0, self._set_progress, 15, "Verifying document type…")
                applicable, reason = _check_loan_applicability(self._extracted_text)
            except Exception as e:
                # If the check itself errors, be lenient and proceed
                applicable, reason = True, ""

            if not applicable:
                msg = (
                    "❌  NOT APPLICABLE FOR LOAN ANALYSIS\n"
                    + "─" * 56 + "\n\n"
                    f"This document cannot be used for loan analysis because\n"
                    f"{reason}\n\n"
                    "Please upload a loan application document such as:\n"
                    "  •  Loan application form (any BSV product)\n"
                    "  •  CI/BI (Credit & Background Investigation) form\n"
                    "  •  Cashflow / income statement of the borrower\n"
                    "  •  Credit approval sheet or financial statement\n"
                    "  •  Any document that is part of a loan application package"
                )
                self.after(0, self._set_progress, 100, "Done")
                self.after(0, self._finish_analysis_error, msg)
                return

            # ── Step 2: full credit analysis ──────────────────────────────
            self.after(0, self._set_progress, 25, "Analysing financials…")
            try:
                result = _run_credit_analysis(self._extracted_text)
            except Exception as e:
                result = f"⚠ Analysis error:\n{type(e).__name__}: {e}"
            self.after(0, self._set_progress, 90, "Formatting report…")
            self.after(0, self._finish_analysis, result)

        threading.Thread(target=worker, daemon=True).start()

    def _finish_analysis(self, result):
        self._show_loader(False)
        self._write_analysis(result, TXT_NAVY)
        self._status_lbl.config(text="●  Analysis Done", fg=LIME_DARK)
        self._analyze_btn.configure(state="normal")

    def _finish_analysis_error(self, msg):
        """Called when the document fails the applicability check."""
        self._show_loader(False)
        self._write_analysis(msg, ACCENT_RED)
        self._status_lbl.config(text="●  Not Applicable", fg=ACCENT_RED)
        self._analyze_btn.configure(state="normal")

    # ── SEARCH ────────────────────────────────────────────────────────────
    def _active_textbox(self):
        """Return the Text widget for whichever tab is visible."""
        return self._textbox if self._current_tab == "extract" else self._analysis_box

    def _do_search(self, *_):
        query = self._search_var.get().strip()
        box   = self._active_textbox()

        # Clear previous highlights
        box.tag_remove("search_match",   "1.0", "end")
        box.tag_remove("search_current", "1.0", "end")
        self._search_matches = []
        self._search_cursor  = -1

        if not query:
            self._match_lbl.config(text="", fg=TXT_SOFT)
            return

        # Find all matches (case-insensitive)
        start = "1.0"
        while True:
            pos = box.search(query, start, stopindex="end", nocase=True)
            if not pos:
                break
            end = f"{pos}+{len(query)}c"
            self._search_matches.append((pos, end))
            box.tag_add("search_match", pos, end)
            start = end

        count = len(self._search_matches)
        if count == 0:
            self._match_lbl.config(text="No results", fg=ACCENT_RED)
            return

        # Jump to first match
        self._search_cursor = 0
        self._highlight_current()
        self._match_lbl.config(
            text=f"1 / {count}", fg=LIME_DARK
        )

    def _highlight_current(self):
        if not self._search_matches:
            return
        box = self._active_textbox()
        # Remove old current highlight, re-apply all matches, then mark current
        box.tag_remove("search_current", "1.0", "end")
        pos, end = self._search_matches[self._search_cursor]
        box.tag_add("search_current", pos, end)
        box.see(pos)
        total = len(self._search_matches)
        self._match_lbl.config(
            text=f"{self._search_cursor + 1} / {total}", fg=LIME_DARK
        )

    def _search_next(self):
        if not self._search_matches:
            return
        self._search_cursor = (self._search_cursor + 1) % len(self._search_matches)
        self._highlight_current()

    def _search_prev(self):
        if not self._search_matches:
            return
        self._search_cursor = (self._search_cursor - 1) % len(self._search_matches)
        self._highlight_current()

    def _clear_search(self):
        self._search_var.set("")
        box = self._active_textbox()
        box.tag_remove("search_match",   "1.0", "end")
        box.tag_remove("search_current", "1.0", "end")
        self._search_matches = []
        self._search_cursor  = -1
        self._match_lbl.config(text="", fg=TXT_SOFT)
        self._search_entry.focus_set()

    def _copy(self):
        if self._current_tab == "extract":
            content = self._textbox.get("1.0","end").strip()
        else:
            content = self._analysis_box.get("1.0","end").strip()

        skip = ("Results will appear here", "Loan analysis will appear here")
        if content and not any(s in content for s in skip):
            self.clipboard_clear()
            self.clipboard_append(content)
            self._copy_btn.configure(text="✓  Copied!")
            self.after(2200, lambda: self._copy_btn.configure(text="⎘  Copy All"))


# ── Entry ─────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    app = DocExtractorApp()
    app.mainloop()