"""
file_generator.py — DocExtract Pro
====================================
Detects when the AI chatbot response should be saved as a file,
then generates DOCX / XLSX / PDF using pure Python libraries.

Supported triggers (detected in the user's message):
  • "generate a word / docx document"
  • "create an excel / xlsx / spreadsheet"
  • "export as pdf / generate a pdf"
  • "save as a file / download as …"

Public API
----------
    detect_file_intent(user_message: str) -> str | None
        Returns "docx" | "xlsx" | "pdf" | None

    generate_file(ai_reply: str, file_type: str,
                  doc_text: str = "",
                  suggested_name: str = "output") -> Path
        Writes the file to a temp location and returns its Path.
"""

import re
import os
import tempfile
from pathlib import Path
from datetime import datetime

# ── Intent keywords ────────────────────────────────────────────────────────────

_DOCX_KEYWORDS = {
    "word document", "word doc", ".docx", "docx file",
    "generate a document", "create a document", "write a document",
    "save as word", "export as word", "download as word",
    "generate a report", "create a report", "write a report",
    "make a document", "produce a document",
}

_XLSX_KEYWORDS = {
    "excel", "spreadsheet", ".xlsx", "xlsx file",
    "excel file", "create a table", "generate a table",
    "save as excel", "export as excel", "download as excel",
    "make a spreadsheet", "create a worksheet file",
    "generate a worksheet file",
}

_PDF_KEYWORDS = {
    ".pdf", "pdf file", "pdf document",
    "save as pdf", "export as pdf", "download as pdf",
    "generate a pdf", "create a pdf", "make a pdf",
}


def detect_file_intent(user_message: str) -> str | None:
    """
    Inspect *user_message* and return the file type the user wants,
    or None if no file generation is requested.

    Priority: docx > xlsx > pdf
    """
    q = user_message.lower()
    for kw in _DOCX_KEYWORDS:
        if kw in q:
            return "docx"
    for kw in _XLSX_KEYWORDS:
        if kw in q:
            return "xlsx"
    for kw in _PDF_KEYWORDS:
        if kw in q:
            return "pdf"
    return None


# ── Output directory ───────────────────────────────────────────────────────────

def _output_dir() -> Path:
    """Return a writable directory for generated files (Desktop preferred)."""
    desktop = Path.home() / "Desktop"
    if desktop.exists():
        out = desktop / "DocExtract_Files"
    else:
        out = Path(tempfile.gettempdir()) / "DocExtract_Files"
    out.mkdir(parents=True, exist_ok=True)
    return out


def _unique_path(stem: str, ext: str) -> Path:
    ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
    name = f"{stem}_{ts}.{ext}"
    return _output_dir() / name


# ══════════════════════════════════════════════════════════════════════════════
#  DOCX GENERATOR  (python-docx)
# ══════════════════════════════════════════════════════════════════════════════

def _generate_docx(ai_reply: str, doc_text: str, stem: str) -> Path:
    """
    Convert *ai_reply* (markdown-ish text) into a formatted .docx file.
    Falls back to a plain-text dump if python-docx is not installed.
    """
    out = _unique_path(stem, "docx")

    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # ── Page margins (1 inch all around) ─────────────────────────────
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1)
            section.right_margin  = Inches(1)

        # ── Default style ──────────────────────────────────────────────────
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        # ── Header ────────────────────────────────────────────────────────
        hdr = doc.add_heading("Banco San Vicente — AI Analysis Report", level=1)
        hdr.runs[0].font.color.rgb = RGBColor(0x1C, 0x2E, 0x7A)  # NAVY_DEEP

        ts_para = doc.add_paragraph(
            f"Generated: {datetime.now().strftime('%B %d, %Y  %H:%M')}"
        )
        ts_para.runs[0].font.color.rgb = RGBColor(0x9A, 0xAA, 0xCE)  # TXT_MUTED
        ts_para.runs[0].font.size = Pt(9)

        doc.add_paragraph()  # spacer

        # ── AI reply — parse headings / bullets / body ─────────────────────
        for line in ai_reply.splitlines():
            s = line.rstrip()

            if re.match(r'^#{1,2}\s+', s):
                clean = re.sub(r'^#{1,2}\s+', '', s).strip()
                h = doc.add_heading(clean, level=2)
                h.runs[0].font.color.rgb = RGBColor(0x2B, 0x45, 0xA8)
            elif re.match(r'^#{3,4}\s+', s):
                clean = re.sub(r'^#{3,4}\s+', '', s).strip()
                h = doc.add_heading(clean, level=3)
            elif re.match(r'^\d{1,2}\.\s+[A-Z]', s):
                h = doc.add_heading(s, level=2)
                h.runs[0].font.color.rgb = RGBColor(0x1C, 0x2E, 0x7A)
            elif re.match(r'^\s*[•\-\*–]\s', s):
                clean = re.sub(r'^\s*[•\-\*–]\s*', '', s).strip()
                p     = doc.add_paragraph(style="List Bullet")
                run   = p.add_run(clean)
                run.font.size = Pt(11)
            elif set(s.strip()).issubset(set("─—-=_*")) and len(s.strip()) > 4:
                doc.add_paragraph("─" * 60)
            elif s.strip() == "":
                doc.add_paragraph()
            else:
                # Inline bold: **text**
                para = doc.add_paragraph()
                parts = re.split(r'(\*\*[^*]+\*\*)', s)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = para.add_run(part[2:-2])
                        run.bold = True
                    else:
                        para.add_run(part)

        # ── Optional: source document snippet ──────────────────────────────
        if doc_text.strip():
            doc.add_page_break()
            doc.add_heading("Source Document Extract (First 2,000 chars)", level=2)
            snippet = doc_text[:2_000]
            if len(doc_text) > 2_000:
                snippet += "\n\n[… truncated …]"
            p = doc.add_paragraph(snippet)
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0x6A, 0x78, 0xB8)

        doc.save(str(out))
        return out

    except ImportError:
        # Fallback: plain UTF-8 text saved with .docx extension so it opens
        out.write_text(ai_reply, encoding="utf-8")
        return out


# ══════════════════════════════════════════════════════════════════════════════
#  XLSX GENERATOR  (openpyxl)
# ══════════════════════════════════════════════════════════════════════════════

def _generate_xlsx(ai_reply: str, doc_text: str, stem: str) -> Path:
    """
    Write the AI reply into an Excel workbook.
    Tables (Markdown pipe syntax) are detected and formatted as proper cells.
    Plain text goes into a "Report" sheet, raw doc text to a "Source" sheet.
    """
    out = _unique_path(stem, "xlsx")

    try:
        import openpyxl
        from openpyxl.styles import (
            Font, PatternFill, Alignment, Border, Side
        )

        NAVY   = "1C2E7A"
        LIME   = "A8D818"
        GHOST  = "C8D4F8"
        WHITE  = "FFFFFF"
        MIST   = "F0F3FF"

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "AI Analysis"

        # ── Column widths ───────────────────────────────────────────────
        ws.column_dimensions["A"].width = 60
        ws.column_dimensions["B"].width = 30

        # ── Title row ───────────────────────────────────────────────────
        ws["A1"] = "Banco San Vicente — AI Analysis Report"
        ws["A1"].font      = Font(name="Arial", bold=True, size=14, color=WHITE)
        ws["A1"].fill      = PatternFill("solid", fgColor=NAVY)
        ws["A1"].alignment = Alignment(horizontal="left", vertical="center",
                                       wrap_text=True)
        ws.row_dimensions[1].height = 28
        ws.merge_cells("A1:B1")

        ws["A2"] = f"Generated: {datetime.now().strftime('%B %d, %Y  %H:%M')}"
        ws["A2"].font      = Font(name="Arial", size=9, color="9AAACE")
        ws["A2"].alignment = Alignment(horizontal="left")
        ws.merge_cells("A2:B2")

        thin = Side(style="thin", color="D4DCF8")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        row = 4

        def _header_style(cell, text):
            cell.value     = text
            cell.font      = Font(name="Arial", bold=True, size=11, color=WHITE)
            cell.fill      = PatternFill("solid", fgColor="2B45A8")
            cell.alignment = Alignment(horizontal="left", vertical="center",
                                       wrap_text=True)
            cell.border    = border

        def _body_style(cell, text, is_alt=False):
            cell.value     = text
            cell.font      = Font(name="Arial", size=10, color="1C2E7A")
            cell.fill      = PatternFill("solid", fgColor=(MIST if is_alt else WHITE))
            cell.alignment = Alignment(horizontal="left", vertical="top",
                                       wrap_text=True)
            cell.border    = border

        # ── Parse AI reply line by line ─────────────────────────────────
        alt = False
        for line in ai_reply.splitlines():
            s = line.rstrip()

            # Markdown table row  |col|col|
            if s.startswith("|") and s.endswith("|"):
                cols = [c.strip() for c in s.strip("|").split("|")]
                if all(set(c).issubset(set(":-")) for c in cols):
                    continue  # separator row
                for ci, col_val in enumerate(cols, start=1):
                    cell = ws.cell(row=row, column=ci)
                    if row == row:  # detect if it's the first table row (header)
                        _header_style(cell, col_val)
                ws.row_dimensions[row].height = 18
                row += 1
                continue

            # Heading lines
            if re.match(r'^#{1,2}\s+', s):
                clean = re.sub(r'^#{1,4}\s+', '', s).strip()
                c = ws.cell(row=row, column=1, value=clean)
                c.font      = Font(name="Arial", bold=True, size=12, color=WHITE)
                c.fill      = PatternFill("solid", fgColor=NAVY)
                c.alignment = Alignment(horizontal="left", vertical="center")
                c.border    = border
                ws.merge_cells(
                    start_row=row, start_column=1,
                    end_row=row, end_column=2
                )
                ws.row_dimensions[row].height = 22
                row += 1
                alt = False
                continue

            if re.match(r'^#{3,4}\s+', s):
                clean = re.sub(r'^#{3,4}\s+', '', s).strip()
                c = ws.cell(row=row, column=1, value=clean)
                c.font      = Font(name="Arial", bold=True, size=11, color="2B45A8")
                c.fill      = PatternFill("solid", fgColor=GHOST)
                c.alignment = Alignment(horizontal="left")
                c.border    = border
                ws.merge_cells(
                    start_row=row, start_column=1,
                    end_row=row, end_column=2
                )
                ws.row_dimensions[row].height = 18
                row += 1
                continue

            # Blank line → small spacer
            if s.strip() == "":
                ws.row_dimensions[row].height = 6
                row += 1
                alt = False
                continue

            # Bullet
            if re.match(r'^\s*[•\-\*–]\s', s):
                clean = "  • " + re.sub(r'^\s*[•\-\*–]\s*', '', s).strip()
                c = ws.cell(row=row, column=1, value=clean)
                _body_style(c, clean, alt)
                ws.merge_cells(
                    start_row=row, start_column=1,
                    end_row=row, end_column=2
                )
                ws.row_dimensions[row].height = 16
                row += 1
                alt = not alt
                continue

            # Separator line
            if set(s.strip()).issubset(set("─—-=_")) and len(s.strip()) > 4:
                c = ws.cell(row=row, column=1)
                c.border = Border(bottom=Side(style="medium", color="2B45A8"))
                ws.row_dimensions[row].height = 4
                row += 1
                continue

            # Body text — split on colon for key-value pairs
            if ":" in s and not s.startswith(" "):
                parts = s.split(":", 1)
                key_cell = ws.cell(row=row, column=1, value=parts[0].strip())
                key_cell.font      = Font(name="Arial", bold=True, size=10, color="1C2E7A")
                key_cell.fill      = PatternFill("solid", fgColor=(MIST if alt else WHITE))
                key_cell.border    = border
                key_cell.alignment = Alignment(wrap_text=True, vertical="top")
                ws.column_dimensions["A"].width = max(
                    ws.column_dimensions["A"].width, min(len(parts[0]) + 4, 40)
                )
                val_cell = ws.cell(row=row, column=2, value=parts[1].strip())
                _body_style(val_cell, parts[1].strip(), alt)
            else:
                c = ws.cell(row=row, column=1, value=s)
                _body_style(c, s, alt)
                ws.merge_cells(
                    start_row=row, start_column=1,
                    end_row=row, end_column=2
                )

            ws.row_dimensions[row].height = 16
            row += 1
            alt = not alt

        # ── Source sheet ──────────────────────────────────────────────────
        if doc_text.strip():
            ws2 = wb.create_sheet("Source Document")
            ws2.column_dimensions["A"].width = 100
            ws2["A1"] = "Source Document (First 3,000 chars)"
            ws2["A1"].font = Font(name="Arial", bold=True, size=12, color=WHITE)
            ws2["A1"].fill = PatternFill("solid", fgColor=NAVY)
            ws2["A1"].alignment = Alignment(horizontal="left", vertical="center")
            ws2.row_dimensions[1].height = 22

            snippet = doc_text[:3_000]
            r = 3
            for ln in snippet.splitlines():
                c = ws2.cell(row=r, column=1, value=ln)
                c.font      = Font(name="Consolas", size=9, color="1C2E7A")
                c.alignment = Alignment(wrap_text=False)
                ws2.row_dimensions[r].height = 14
                r += 1

        wb.save(str(out))
        return out

    except ImportError:
        # Fallback: CSV
        out_csv = out.with_suffix(".csv")
        out_csv.write_text(ai_reply, encoding="utf-8")
        return out_csv


# ══════════════════════════════════════════════════════════════════════════════
#  PDF GENERATOR  (reportlab)
# ══════════════════════════════════════════════════════════════════════════════

def _generate_pdf(ai_reply: str, doc_text: str, stem: str) -> Path:
    """
    Render the AI reply as a styled PDF using reportlab.
    Falls back to a plain .txt if reportlab is not installed.
    """
    out = _unique_path(stem, "pdf")

    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer,
            HRFlowable, ListFlowable, ListItem
        )
        from reportlab.lib.enums import TA_LEFT, TA_CENTER

        NAVY_RL   = colors.HexColor("#1C2E7A")
        NAVY_MID  = colors.HexColor("#2B45A8")
        LIME_RL   = colors.HexColor("#A8D818")
        SOFT_RL   = colors.HexColor("#6878B8")
        MUTED_RL  = colors.HexColor("#9AAACE")

        doc = SimpleDocTemplate(
            str(out),
            pagesize=LETTER,
            leftMargin=inch, rightMargin=inch,
            topMargin=inch,  bottomMargin=inch,
        )

        styles = getSampleStyleSheet()

        s_title   = ParagraphStyle("BSVTitle",  parent=styles["Title"],
                                   fontName="Helvetica-Bold", fontSize=16,
                                   textColor=NAVY_RL, spaceAfter=4)
        s_meta    = ParagraphStyle("BSVMeta",   parent=styles["Normal"],
                                   fontName="Helvetica", fontSize=8,
                                   textColor=MUTED_RL, spaceAfter=12)
        s_h1      = ParagraphStyle("BSVH1",     parent=styles["Heading1"],
                                   fontName="Helvetica-Bold", fontSize=13,
                                   textColor=NAVY_MID, spaceBefore=14, spaceAfter=4)
        s_h2      = ParagraphStyle("BSVH2",     parent=styles["Heading2"],
                                   fontName="Helvetica-Bold", fontSize=11,
                                   textColor=NAVY_RL, spaceBefore=10, spaceAfter=3)
        s_body    = ParagraphStyle("BSVBody",   parent=styles["Normal"],
                                   fontName="Helvetica", fontSize=10,
                                   textColor=colors.HexColor("#1C2E7A"),
                                   leading=15, spaceAfter=4)
        s_bullet  = ParagraphStyle("BSVBullet", parent=s_body,
                                   leftIndent=20, bulletIndent=8,
                                   spaceAfter=2)
        s_source  = ParagraphStyle("BSVSource", parent=styles["Normal"],
                                   fontName="Courier", fontSize=8,
                                   textColor=SOFT_RL, leading=11)

        story = []

        story.append(Paragraph("Banco San Vicente — AI Analysis Report", s_title))
        story.append(Paragraph(
            f"Generated: {datetime.now().strftime('%B %d, %Y  %H:%M')}",
            s_meta
        ))
        story.append(HRFlowable(width="100%", thickness=2,
                                color=LIME_RL, spaceAfter=12))

        for line in ai_reply.splitlines():
            s = line.rstrip()

            if re.match(r'^#{1,2}\s+', s):
                clean = re.sub(r'^#{1,2}\s+', '', s).strip()
                story.append(Paragraph(clean, s_h1))
            elif re.match(r'^#{3,4}\s+', s):
                clean = re.sub(r'^#{3,4}\s+', '', s).strip()
                story.append(Paragraph(clean, s_h2))
            elif re.match(r'^\d{1,2}\.\s+[A-Z]', s):
                story.append(Paragraph(s, s_h1))
            elif re.match(r'^\s*[•\-\*–]\s', s):
                clean = re.sub(r'^\s*[•\-\*–]\s*', '', s).strip()
                clean = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', clean)
                story.append(Paragraph(f"• {clean}", s_bullet))
            elif set(s.strip()).issubset(set("─—-=_")) and len(s.strip()) > 4:
                story.append(HRFlowable(width="100%", thickness=1,
                                        color=MUTED_RL, spaceAfter=6))
            elif s.strip() == "":
                story.append(Spacer(1, 6))
            else:
                # Inline bold
                safe = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', s)
                story.append(Paragraph(safe, s_body))

        # Source doc appendix
        if doc_text.strip():
            story.append(Spacer(1, 20))
            story.append(HRFlowable(width="100%", thickness=2,
                                    color=LIME_RL, spaceAfter=8))
            story.append(Paragraph("Source Document (First 2,000 chars)", s_h1))
            snippet = doc_text[:2_000]
            if len(doc_text) > 2_000:
                snippet += "\n[… truncated …]"
            for ln in snippet.splitlines():
                story.append(Paragraph(ln or " ", s_source))

        doc.build(story)
        return out

    except ImportError:
        out_txt = out.with_suffix(".txt")
        out_txt.write_text(ai_reply, encoding="utf-8")
        return out_txt


# ══════════════════════════════════════════════════════════════════════════════
#  PUBLIC ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════

def generate_file(
    ai_reply: str,
    file_type: str,
    doc_text: str = "",
    suggested_name: str = "bsv_analysis",
) -> Path:
    """
    Generate a file of *file_type* ("docx" | "xlsx" | "pdf") from *ai_reply*.

    Parameters
    ----------
    ai_reply       : The full AI response text (may contain Markdown).
    file_type      : One of "docx", "xlsx", "pdf".
    doc_text       : Optional extracted source document text for an appendix.
    suggested_name : Stem for the output filename (timestamp appended).

    Returns
    -------
    Path to the generated file.
    """
    stem = re.sub(r'[^\w\-]', '_', suggested_name)[:40] or "bsv_analysis"

    if file_type == "docx":
        return _generate_docx(ai_reply, doc_text, stem)
    elif file_type == "xlsx":
        return _generate_xlsx(ai_reply, doc_text, stem)
    elif file_type == "pdf":
        return _generate_pdf(ai_reply, doc_text, stem)
    else:
        raise ValueError(f"Unsupported file_type: {file_type!r}")