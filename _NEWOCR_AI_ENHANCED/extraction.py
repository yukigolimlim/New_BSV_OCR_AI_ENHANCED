"""
extraction.py — DocExtract Pro
================================
All document parsing logic.  Each parser accepts a file path and an optional
progress_cb(pct, stage_str) callback so the UI can update its progress bar.

Public API
----------
  extract(file_path, progress_cb=None) -> str
      Dispatch to the correct parser based on file extension.

OPTIMIZATIONS (Phase 1 API call reduction)
-------------------------------------------
  1. Confidence check call REMOVED — saves 1 Gemini call per OCR attempt.
  2. Retries reduced from 3 attempts → 2 (1 retry max).
  3. pdfplumber text-density threshold raised 50 → 150 chars/page.

PP-STRUCTURE PATCH (Phase 2 VLM reduction)
-------------------------------------------
  PP-1. PP-StructureV3 replaces plain PaddleOCR in Stage 2.
  PP-2. Quality gate added after Stage 2.
  PP-3. VLM quota fallback improved.
  PP-4. Fallback model updated: gemini-2.0-flash → gemini-2.5-flash-lite

MERGE PATCH (Phase 3 — accuracy improvement for digital PDFs)
--------------------------------------------------------------
  MG-1. Two-pass merge: pdfplumber text + PP-Structure layout.
  MG-2. Graceful fallback chain within the merge.

IMAGE MERGE PATCH (Phase 4 — same accuracy gains for image files)
------------------------------------------------------------------
  IM-1. Two-pass image merge: PaddleOCR words + PP-Structure layout.
  IM-2. Fallback chain for image merge.

GPU PATCH (Phase 6 — GPU acceleration)
-----------------------------------------
  GPU-1. PaddleOCR now uses GPU by default (PADDLE_USE_GPU defaults to "1").
         Override by setting PADDLE_USE_GPU=0 in the .env file to fall back to CPU.
  GPU-2. PPStructure (_make_pp_structure_engine) now respects the same
         PADDLE_USE_GPU flag, enabling GPU inference for layout detection and
         table recognition in both the main-thread singleton and all
         per-page worker threads.

SPEED PATCH (Phase 5 — parallel processing)
--------------------------------------------
  SP-1  FIX-1: Parallel PDF page processing.
        PP-Structure and pdfplumber word extraction now run concurrently
        across pages using ThreadPoolExecutor.  Each worker thread gets its
        own PPStructure engine instance (thread-local) so the shared singleton
        is never accessed from multiple threads simultaneously — preserving
        full accuracy while cutting multi-page PDF processing time.

  SP-2  FIX-2: Parallel multi-document extraction in CIBI Stage 3.
        extract_multiple_parallel() runs each file through extract() in its
        own thread.  Results are collected in original order.  UI progress
        callbacks are marshalled through self.after(0, ...) as required by
        Tkinter's threading model.

  SP-3  FIX-3: CIC cache reuse across Stage 1 → Stage 4.
        extract_cic_cached() accepts an already-extracted text string and
        returns the cached JSON if available, skipping the Gemini call that
        populate_cibi_form() would otherwise repeat.  Called from app.py
        _cibi_run_stage1() to prime the cache before Stage 4 fires.

  SP-4  FIX-5: TrOCR background prewarming.
        prewarm_trocr() loads microsoft/trocr-base-handwritten in a daemon
        thread.  Called from app.py __init__ so the model is warm before the
        user reaches the Bank CI step.

  SP-5  FIX-7: PaddleOCR + PP-Structure background prewarming.
        prewarm_ocr_engines() loads both singletons in a daemon thread.
        Called from app.py _prewarm_rag() so all local ML models load at
        startup rather than on first use.
"""
from __future__ import annotations
from dataclasses import dataclass, field
from typing import Optional
import os
import sys
import base64
import traceback
import logging
import numpy as np
import re
import re as _re
import tempfile
import threading                          # SP-1 SP-2
from concurrent.futures import (          # SP-1 SP-2
    ThreadPoolExecutor, as_completed
)
from pathlib import Path
from dotenv import load_dotenv
import traceback
import numpy as np
load_dotenv()

from utils import SCRIPT_DIR, IMAGE_EXTS, POPPLER_PATH
from cic_parser import is_cic_report, format_cic_for_analysis

logger = logging.getLogger(__name__)


# ── Lazy singletons ───────────────────────────────────────────────────────────
_gemini_client       = None
_paddle_ocr          = None
_pp_structure_engine = None          # PP-STRUCTURE PATCH PP-1
_trocr_processor     = None          # SP-4 / GPU-3: TrOCR processor singleton
_trocr_model         = None          # SP-4 / GPU-3: TrOCR model singleton

# SP-1: Lock protecting the shared PP-Structure singleton.
# Individual page workers create their own local engine instances (see
# _make_pp_structure_engine()) so this lock is only needed when the global
# singleton is first initialised.
_pp_structure_lock = threading.Lock()
_paddle_ocr_lock   = threading.Lock()
_trocr_lock        = threading.Lock()  # SP-4 / GPU-3


# ── GPU availability helper ────────────────────────────────────────────────────

def _cuda_available() -> bool:
    """Return True if Paddle sees a CUDA-capable GPU."""
    try:
        import paddle
        return paddle.device.cuda.device_count() > 0
    except Exception:
        return False


def _use_gpu() -> bool:
    """
    Master GPU switch used by all local ML engines.

    Priority:
      1. PADDLE_USE_GPU env var explicitly set -> honour it as a boolean.
      2. Otherwise -> default to CPU (no startup CUDA probing/retries).

    This avoids startup lag on machines where CUDA/Torch libraries are
    partially installed or mismatched.
    """
    env_val = os.getenv("PADDLE_USE_GPU", "").strip()
    explicit = (env_val == "1")
    if not explicit:
        return False
    if not _cuda_available():
        logger.warning(
            "PADDLE_USE_GPU=1 requested but no CUDA GPU detected -- falling back to CPU."
        )
        return False
    return True


def _get_gemini():
    global _gemini_client
    if _gemini_client is None:
        from google import genai
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise EnvironmentError(
                "GEMINI_API_KEY not found.\n\n"
                "Create a .env file:\n"
                "    GEMINI_API_KEY=your_key_here\n\n"
                "Get a free key at: https://aistudio.google.com/app/apikey"
            )
        _gemini_client = genai.Client(api_key=api_key)
    return _gemini_client


def _get_paddle_ocr():
    global _paddle_ocr
    if _paddle_ocr is None:
        with _paddle_ocr_lock:
            if _paddle_ocr is None:          # double-checked locking
                from paddleocr import PaddleOCR
                _paddle_ocr = PaddleOCR(
                    use_angle_cls=True, lang="en", use_gpu=_use_gpu()
                )
    return _paddle_ocr


# ── PP-STRUCTURE PATCH PP-1 + SP-1 ───────────────────────────────────────────

def _make_pp_structure_engine():
    """
    SP-1: Create a fresh PPStructure engine instance.
    Used by parallel page workers so each thread has its own engine —
    PPStructure is NOT thread-safe and must not be shared across threads.
    """
    from paddleocr import PPStructure
    return PPStructure(
        show_log=False,
        lang="en",
        return_ocr_result_in_table=True,
        use_gpu=_use_gpu(),
    )


def _get_pp_structure():
    """
    Lazy singleton for the MAIN-THREAD PPStructure engine.
    Never call this from a worker thread — use _make_pp_structure_engine()
    inside the worker instead.
    """
    global _pp_structure_engine
    if _pp_structure_engine is None:
        with _pp_structure_lock:
            if _pp_structure_engine is None:   # double-checked locking
                _pp_structure_engine = _make_pp_structure_engine()
    return _pp_structure_engine


# ══════════════════════════════════════════════════════════════════════════════
#  SP-4 / GPU-3: TrOCR lazy singleton
# ══════════════════════════════════════════════════════════════════════════════

def _get_trocr():
    """
    Lazy singleton loader for microsoft/trocr-base-handwritten.

    Returns (processor, model) tuple.  Both are cached after first load.
    The model is moved to CUDA if a GPU is available, otherwise stays on CPU.

    Thread-safe via _trocr_lock (double-checked locking).
    """
    global _trocr_processor, _trocr_model
    if _trocr_model is None:
        with _trocr_lock:
            if _trocr_model is None:
                from transformers import TrOCRProcessor, VisionEncoderDecoderModel
                import torch

                model_name = "microsoft/trocr-base-handwritten"
                logger.info("Loading TrOCR model: %s", model_name)

                processor = TrOCRProcessor.from_pretrained(model_name)
                model     = VisionEncoderDecoderModel.from_pretrained(model_name)

                if _use_gpu():
                    device = torch.device("cuda")
                    model  = model.to(device)
                    logger.info("TrOCR loaded on GPU (cuda).")
                else:
                    logger.info("TrOCR loaded on CPU.")

                model.eval()
                _trocr_processor = processor
                _trocr_model     = model

    return _trocr_processor, _trocr_model


# ══════════════════════════════════════════════════════════════════════════════
#  SP-4 FIX-5: TrOCR background prewarming
# ══════════════════════════════════════════════════════════════════════════════

def prewarm_trocr() -> None:
    """
    FIX-5: Load microsoft/trocr-base-handwritten in a background daemon
    thread so it is warm before the user reaches the Bank CI step.

    Call this from app.py __init__ after _build_ui():
        threading.Thread(target=prewarm_trocr, daemon=True).start()

    Safe to call multiple times — subsequent calls are no-ops once loaded.
    """
    def _load():
        try:
            _get_trocr()
            logger.info("TrOCR prewarmed successfully.")
        except Exception as e:
            logger.warning("TrOCR prewarm failed (non-fatal): %s", e)

    t = threading.Thread(target=_load, daemon=True, name="trocr-prewarm")
    t.start()


# ══════════════════════════════════════════════════════════════════════════════
#  SP-5 FIX-7: PaddleOCR + PP-Structure background prewarming
# ══════════════════════════════════════════════════════════════════════════════

def prewarm_ocr_engines() -> None:
    """
    FIX-7: Load PaddleOCR and PPStructure singletons in a background daemon
    thread so the first-use penalty (~3-5 s) is paid at app startup rather
    than on the user's first extraction.

    Call this from app.py _prewarm_rag() alongside the RAG warmup:
        from extraction import prewarm_ocr_engines
        prewarm_ocr_engines()

    Safe to call multiple times — subsequent calls are no-ops once loaded.
    """
    def _load():
        try:
            _get_paddle_ocr()
            logger.info("PaddleOCR prewarmed successfully.")
        except Exception as e:
            logger.warning("PaddleOCR prewarm failed (non-fatal): %s", e)
        try:
            _get_pp_structure()
            logger.info("PP-Structure prewarmed successfully.")
        except Exception as e:
            logger.warning("PP-Structure prewarm failed (non-fatal): %s", e)

    t = threading.Thread(target=_load, daemon=True, name="ocr-prewarm")
    t.start()


# ══════════════════════════════════════════════════════════════════════════════
#  SP-2 FIX-2: Parallel multi-document extraction
# ══════════════════════════════════════════════════════════════════════════════

def extract_parallel(
    file_paths:  list[str],
    progress_cb=None,
    max_workers: int = 4,
) -> list[str]:
    """
    FIX-2: Extract multiple files concurrently using a thread pool.

    Each file is processed by extract() in its own worker thread.
    Results are returned in the SAME ORDER as file_paths regardless of
    which file finishes first.

    progress_cb is called with (pct, message) from the main thread via
    a thread-safe wrapper — never from worker threads directly, which
    would crash Tkinter.

    Parameters
    ----------
    file_paths  : list of absolute file paths to extract
    progress_cb : optional (pct: int, stage: str) -> None callback
    max_workers : thread pool size (default 4; BSV CIBI typically has 3 docs)

    Returns
    -------
    list of extracted text strings, same length and order as file_paths.
    Empty string for any file that raised an exception.
    """
    if not file_paths:
        return []
    if len(file_paths) == 1:
        # No benefit to threading a single file
        return [extract(file_paths[0], progress_cb)]

    total   = len(file_paths)
    results = [""] * total

    # Thread-safe progress wrapper — stores messages; caller polls or uses
    # after(0, ...) to update UI.
    _lock         = threading.Lock()
    _done_count   = 0
    _status: dict = {}   # index → latest stage message

    def _cb_for(idx: int, name: str):
        """Return a per-file progress callback that is thread-safe."""
        def _inner(pct: int, stage: str = ""):
            nonlocal _done_count
            with _lock:
                _status[idx] = f"[{name}] {stage or str(pct) + '%'}"
                combined_pct = int(
                    sum(
                        pct if i == idx else (100 if results[i] else 0)
                        for i in range(total)
                    ) / total
                )
            if progress_cb:
                try:
                    progress_cb(
                        combined_pct,
                        f"[{idx+1}/{total}] {name}: {stage or str(pct)+'%'}"
                    )
                except Exception:
                    pass
        return _inner

    def _worker(idx: int, fp: str):
        name = Path(fp).name
        try:
            text = extract(fp, _cb_for(idx, name))
        except Exception as e:
            logger.error(
                "extract_parallel: file %d (%s) raised: %s", idx, name, e
            )
            text = f"[Extraction error: {e}]"
        results[idx] = text
        return idx

    with ThreadPoolExecutor(
        max_workers=min(max_workers, total),
        thread_name_prefix="bsv-extract"
    ) as pool:
        futures = {
            pool.submit(_worker, i, fp): i
            for i, fp in enumerate(file_paths)
        }
        for future in as_completed(futures):
            try:
                future.result()
            except Exception as e:
                idx = futures[future]
                logger.error(
                    "extract_parallel future %d raised: %s", idx, e
                )

    if progress_cb:
        try:
            progress_cb(100, "All files extracted.")
        except Exception:
            pass

    return results


# ══════════════════════════════════════════════════════════════════════════════
#  SP-3 FIX-3: CIC cache-priming helper
# ══════════════════════════════════════════════════════════════════════════════

def prime_cic_cache(cic_text: str, api_key: str) -> None:
    """
    FIX-3: Trigger the CIC Gemini extraction NOW (during Stage 1) so that
    when populate_cibi_form() calls extract_cic() in Stage 4, it finds a
    warm cache entry and skips the API call entirely.

    Call this from app.py _cibi_run_stage1() worker thread, right after the
    raw CIC text has been extracted from the file:

        if cic_text:
            from extraction import prime_cic_cache
            prime_cic_cache(cic_text, GEMINI_API_KEY)

    This runs synchronously in the Stage 1 worker thread — no new thread
    needed because _cibi_run_stage1() already runs off the main thread.

    Safe to call even if cic_text is empty (no-op).
    """
    if not cic_text or not cic_text.strip():
        return
    try:
        from Cibi_populator import extract_cic
        logger.info(
            "FIX-3: Priming CIC cache (%d chars) during Stage 1…",
            len(cic_text)
        )
        extract_cic(cic_text, api_key)
        logger.info("FIX-3: CIC cache primed — Stage 4 will be a cache hit.")
    except Exception as e:
        # Non-fatal — Stage 4 will just re-extract normally
        logger.warning("FIX-3: CIC cache priming failed (non-fatal): %s", e)


# ── PP-STRUCTURE helpers (unchanged from original) ───────────────────────────

def _html_table_to_text(html: str) -> str:
    rows  = re.findall(r"<tr[^>]*>(.*?)</tr>", html, re.S)
    lines = []
    for row in rows:
        cells = re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", row, re.S)
        cells = [re.sub(r"<[^>]+>", "", c).strip() for c in cells]
        cells = [c for c in cells if c]
        if cells:
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def _pp_structure_extract(img_array, engine=None) -> str:
    """
    Run PP-StructureV3 on a preprocessed numpy image array.

    SP-1: Accepts an optional `engine` parameter.  Worker threads pass their
    own local engine instance; the main thread uses the shared singleton.
    """
    if engine is None:
        engine = _get_pp_structure()
    result = engine(img_array)
    parts: list[str] = []

    for region in result:
        region_type = region.get("type", "").lower()
        res         = region.get("res", None)

        if region_type == "table" and isinstance(res, dict):
            html = res.get("html", "")
            if html:
                table_text = _html_table_to_text(html)
                if table_text.strip():
                    parts.append(f"[TABLE]\n{table_text}\n[/TABLE]")

        elif region_type in ("text", "title", "header",
                             "footer", "reference", "figure_caption"):
            if isinstance(res, list):
                lines = []
                for item in res:
                    if not item:
                        continue
                    if isinstance(item, dict):
                        t    = item.get("text", "") or ""
                        conf = float(item.get("confidence", 1.0))
                        if t.strip() and conf > 0.5:
                            lines.append(t.strip())
                    elif (isinstance(item, (list, tuple))
                          and len(item) > 1
                          and isinstance(item[1], (list, tuple))
                          and len(item[1]) >= 2):
                        t    = item[1][0] or ""
                        conf = float(item[1][1])
                        if t.strip() and conf > 0.5:
                            lines.append(t.strip())
                if lines:
                    parts.append("\n".join(lines))

    return "\n\n".join(parts)


# ── PP-STRUCTURE PATCH PP-2: quality gate helper ──────────────────────────────

def _text_quality_score(text: str) -> float:
    if not text or len(text.strip()) < 100:
        return 0.0
    readable = sum(
        1 for c in text
        if c.isalnum() or c in " .,₱/-:()\n\t%@#'|[]"
    )
    return readable / max(len(text), 1)


# ── MERGE PATCH MG-1: bbox helpers ───────────────────────────────────────────

def _bbox_overlap(b1: tuple, b2: tuple) -> float:
    ix0 = max(b1[0], b2[0])
    iy0 = max(b1[1], b2[1])
    ix1 = min(b1[2], b2[2])
    iy1 = min(b1[3], b2[3])
    if ix1 <= ix0 or iy1 <= iy0:
        return 0.0
    return (ix1 - ix0) * (iy1 - iy0)


def _words_in_bbox(
    words:     list,
    bbox:      tuple,
    threshold: float = 0.3,
) -> str:
    matched = []
    for w in words:
        wb        = (w["x0"], w["top"], w["x1"], w["bottom"])
        word_area = max((wb[2] - wb[0]) * (wb[3] - wb[1]), 1)
        if _bbox_overlap(wb, bbox) / word_area >= threshold:
            matched.append(w)
    matched.sort(key=lambda w: (round(w["top"] / 5) * 5, w["x0"]))
    return " ".join(w["text"] for w in matched)


def _extract_table_with_plumber_text(
    html:            str,
    pl_words:        list,
    region_bbox:     list,
    scale_x:         float,
    scale_y:         float,
    words_in_tables: set,
) -> list:
    rows_html = re.findall(r"<tr[^>]*>(.*?)</tr>", html, re.S)
    if not rows_html:
        return []

    result_rows: list = []
    rx0, ry0, rx1, ry1 = region_bbox
    num_rows            = len(rows_html)

    for row_idx, row_html in enumerate(rows_html):
        cells_html = re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", row_html, re.S)
        if not cells_html:
            continue

        num_cols   = len(cells_html)
        cell_texts: list = []

        for col_idx, cell_html in enumerate(cells_html):
            cell_x0 = rx0 + (rx1 - rx0) * (col_idx     / num_cols)
            cell_x1 = rx0 + (rx1 - rx0) * ((col_idx+1) / num_cols)
            cell_y0 = ry0 + (ry1 - ry0) * (row_idx     / num_rows)
            cell_y1 = ry0 + (ry1 - ry0) * ((row_idx+1) / num_rows)

            pdf_bbox = (
                cell_x0 / scale_x,
                cell_y0 / scale_y,
                cell_x1 / scale_x,
                cell_y1 / scale_y,
            )

            cell_text = _words_in_bbox(pl_words, pdf_bbox)

            for i, w in enumerate(pl_words):
                wb        = (w["x0"], w["top"], w["x1"], w["bottom"])
                word_area = max((wb[2]-wb[0]) * (wb[3]-wb[1]), 1)
                if _bbox_overlap(wb, pdf_bbox) / word_area >= 0.3:
                    words_in_tables.add(i)

            if not cell_text.strip():
                cell_text = re.sub(r"<[^>]+>", "", cell_html).strip()

            cell_texts.append(cell_text)

        row_str = " | ".join(cell_texts)
        if row_str.strip():
            result_rows.append(row_str)

    return result_rows


def _merge_pdfplumber_with_ppstructure(
    pdf_path:    str,
    dpi:         int  = 200,
    progress_cb       = None,
) -> str:
    """
    SP-1 FIX-1: Two-pass extraction for digital PDFs with parallel pages.

    Each page is now processed in its own worker thread via ThreadPoolExecutor.
    Each worker creates its own PPStructure engine instance via
    _make_pp_structure_engine() — the shared singleton (_get_pp_structure())
    is NEVER called from a worker thread, ensuring thread safety without
    sacrificing accuracy.

    pdfplumber word extraction (I/O-bound) overlaps with PP-Structure
    inference (CPU-bound) across pages, cutting multi-page PDF time
    significantly.

    Fallback chain (MG-2) is preserved — per-page errors are isolated and
    never abort the whole document.
    """
    import pdfplumber
    import numpy as np

    kw = ({"poppler_path": POPPLER_PATH}
          if sys.platform == "win32" and Path(POPPLER_PATH).exists() else {})

    # Render all pages up-front
    pdf_images: list = []
    try:
        from pdf2image import convert_from_path
        pdf_images = convert_from_path(
            pdf_path, dpi=dpi, fmt="jpeg",
            jpegopt={"quality": 95, "optimize": True},
            **kw
        )
    except Exception as e:
        logger.warning(
            "pdf2image unavailable in merge pass — "
            "table structure detection skipped: %s", e
        )

    plumber_pdf = None
    try:
        plumber_pdf = pdfplumber.open(pdf_path)
    except Exception as e:
        logger.warning(
            "pdfplumber failed in merge — falling back to PP-Structure only: %s", e
        )
        # Fallback: PP-Structure only, still parallel
        pages_text: list = [""] * len(pdf_images)

        def _pp_only(args):
            i, img = args
            try:
                engine = _make_pp_structure_engine()  # thread-local
                text   = _pp_structure_extract(np.array(img), engine=engine)
                if text.strip():
                    return i, f"=== PAGE {i+1} ===\n{text}"
            except Exception:
                pass
            return i, ""

        with ThreadPoolExecutor(
            max_workers=min(4, len(pdf_images)),
            thread_name_prefix="bsv-pp-only"
        ) as pool:
            for idx, page_text in pool.map(_pp_only,
                                           enumerate(pdf_images)):
                pages_text[idx] = page_text

        return "\n\n".join(p for p in pages_text if p)

    try:
        total_pages = len(plumber_pdf.pages)

        # ── SP-1: Build per-page work items ──────────────────────────────
        page_args = []
        for page_idx in range(total_pages):
            img = pdf_images[page_idx] if page_idx < len(pdf_images) else None
            page_args.append((page_idx, img, plumber_pdf.pages[page_idx]))

        # Ordered results placeholder
        page_results: list[str] = [""] * total_pages

        def _process_page(args):
            """
            SP-1: Worker function — runs in its own thread.
            Creates a THREAD-LOCAL PPStructure engine so the shared
            singleton is never accessed concurrently.
            """
            page_idx, img, pl_page = args

            pl_words:   list = []
            pl_page_w        = 0.0
            pl_page_h        = 0.0
            pl_fulltext      = ""

            # Pass 1: pdfplumber word extraction
            try:
                pl_words    = pl_page.extract_words(
                    x_tolerance      = 3,
                    y_tolerance      = 3,
                    keep_blank_chars = False,
                ) or []
                pl_page_w   = float(pl_page.width)
                pl_page_h   = float(pl_page.height)
                pl_fulltext = pl_page.extract_text() or ""
            except Exception as e:
                logger.warning(
                    "pdfplumber word extraction failed page %d: %s",
                    page_idx + 1, e
                )

            # Pass 2: PP-Structure layout detection
            # SP-1: thread-local engine — never uses shared singleton
            pp_result: list = []
            img_w, img_h    = 0, 0

            if img is not None:
                try:
                    local_engine = _make_pp_structure_engine()
                    img_w, img_h = img.size
                    pp_result    = local_engine(np.array(img))
                except Exception as e:
                    logger.warning(
                        "PP-Structure failed page %d: %s — "
                        "using pdfplumber text only for this page.",
                        page_idx + 1, e
                    )

            scale_x = (img_w / pl_page_w) if (pl_page_w and img_w) else 1.0
            scale_y = (img_h / pl_page_h) if (pl_page_h and img_h) else 1.0

            words_in_tables: set  = set()
            table_blocks:    list = []

            for region in pp_result:
                region_type = region.get("type", "").lower()
                res         = region.get("res", None)
                region_bbox = region.get("bbox", None)

                if region_type == "table" and isinstance(res, dict):
                    html = res.get("html", "")
                    if not html:
                        continue

                    if pl_words and region_bbox:
                        rows = _extract_table_with_plumber_text(
                            html, pl_words, region_bbox,
                            scale_x, scale_y, words_in_tables,
                        )
                        if rows:
                            table_blocks.append(
                                "[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]"
                            )
                    else:
                        table_text = _html_table_to_text(html)
                        if table_text.strip():
                            table_blocks.append(
                                f"[TABLE]\n{table_text}\n[/TABLE]"
                            )

            page_parts: list = []

            if pl_words:
                remaining = [
                    w for i, w in enumerate(pl_words)
                    if i not in words_in_tables
                ]
                if remaining:
                    remaining.sort(
                        key=lambda w: (round(w["top"] / 10) * 10, w["x0"])
                    )
                    non_table = " ".join(w["text"] for w in remaining)
                    if non_table.strip():
                        page_parts.append(non_table)
            elif pl_fulltext.strip():
                page_parts.append(pl_fulltext)

            page_parts.extend(table_blocks)

            if page_parts:
                return (
                    page_idx,
                    f"=== PAGE {page_idx+1} ===\n" +
                    "\n\n".join(page_parts)
                )
            return page_idx, ""

        # ── SP-1: Run pages in parallel ───────────────────────────────────
        max_page_workers = min(4, total_pages)
        logger.info(
            "SP-1: Processing %d PDF pages with %d workers",
            total_pages, max_page_workers
        )

        with ThreadPoolExecutor(
            max_workers=max_page_workers,
            thread_name_prefix="bsv-pdf-page"
        ) as pool:
            futures = {
                pool.submit(_process_page, args): args[0]
                for args in page_args
            }
            completed = 0
            for future in as_completed(futures):
                completed += 1
                if progress_cb:
                    progress_cb(
                        int(20 + (completed / max(total_pages, 1)) * 70),
                        f"Merging page {completed}/{total_pages}..."
                    )
                try:
                    page_idx, text = future.result()
                    page_results[page_idx] = text
                except Exception as e:
                    page_idx = futures[future]
                    logger.warning(
                        "Page %d worker raised: %s", page_idx + 1, e
                    )

    finally:
        try:
            plumber_pdf.close()
        except Exception:
            pass

    return "\n\n".join(p for p in page_results if p)


# ══════════════════════════════════════════════════════════════════════════════
#  IMAGE PREPROCESSOR  (unchanged)
# ══════════════════════════════════════════════════════════════════════════════
def _preprocess_image(
    file_path:  str,
    sharpness:  float = 2.0,
    contrast:   float = 1.5,
    brightness: float = 1.1,
    dpi_scale:  float = 1.0,
) -> tuple[str, str]:
    from PIL import Image, ImageEnhance, ImageFilter
    import cv2
    import numpy as np

    img = Image.open(file_path).convert("RGB")
    w, h = img.size

    target_w = max(1800, int(w * dpi_scale))
    if w < target_w:
        scale = target_w / w
        img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

    img = ImageEnhance.Sharpness(img).enhance(sharpness)
    img = ImageEnhance.Contrast(img).enhance(contrast)
    img = ImageEnhance.Brightness(img).enhance(brightness)
    img = img.filter(ImageFilter.MedianFilter(size=3))

    try:
        arr  = np.array(img)
        gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
        _, binary = cv2.threshold(
            gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU
        )
        coords = np.column_stack(np.where(binary > 0))
        if len(coords) > 100:
            angle = cv2.minAreaRect(coords.astype(np.float32))[-1]
            if angle < -45:
                angle = 90 + angle
            if abs(angle) > 0.5:
                (rh, rw) = gray.shape
                center   = (rw // 2, rh // 2)
                M        = cv2.getRotationMatrix2D(center, angle, 1.0)
                rotated  = cv2.warpAffine(
                    arr, M, (rw, rh),
                    flags=cv2.INTER_CUBIC,
                    borderMode=cv2.BORDER_REPLICATE
                )
                img = Image.fromarray(rotated)
    except Exception:
        logger.warning("Deskew failed — continuing without rotation correction",
                       exc_info=True)

    fd, tmp_color = tempfile.mkstemp(suffix=".jpg")
    os.close(fd)
    img.save(tmp_color, "JPEG", quality=95)

    tmp_gray = tmp_color
    try:
        arr    = np.array(img)
        gray   = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
        binary = cv2.adaptiveThreshold(
            gray, 255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 10
        )
        binary = cv2.medianBlur(binary, 3)
        fd, tmp_gray = tempfile.mkstemp(suffix=".jpg")
        os.close(fd)
        cv2.imwrite(tmp_gray, binary)
    except Exception:
        logger.warning(
            "Grayscale preprocessing failed — using colour image for PaddleOCR",
            exc_info=True
        )
        tmp_gray = tmp_color

    return tmp_color, tmp_gray


# ══════════════════════════════════════════════════════════════════════════════
#  RETRY CONFIGURATION  (unchanged)
# ══════════════════════════════════════════════════════════════════════════════

_RETRY_STRATEGIES = [
    {
        "name":       "High contrast",
        "sharpness":  3.0,
        "contrast":   2.0,
        "brightness": 1.0,
        "dpi_scale":  1.0,
    },
]

_MIN_RESPONSE_CHARS = 200


# ── IMAGE MERGE PATCH (unchanged logic, SP-1 engine fix applied) ──────────────

def _paddle_words_to_flat(paddle_result: list) -> list:
    words = []
    if not paddle_result or not paddle_result[0]:
        return words
    for item in paddle_result[0]:
        try:
            pts  = item[0]
            text = item[1][0]
            conf = float(item[1][1])
            if not text or not text.strip() or conf < 0.3:
                continue
            xs = [p[0] for p in pts]
            ys = [p[1] for p in pts]
            words.append({
                "x0":     min(xs),
                "top":    min(ys),
                "x1":     max(xs),
                "bottom": max(ys),
                "text":   text.strip(),
            })
        except (IndexError, KeyError, TypeError):
            continue
    return words


def _merge_paddle_with_ppstructure(img_array) -> str:
    """
    IMAGE MERGE PATCH: two-pass extraction for image files and scanned pages.
    SP-1: Uses _get_pp_structure() (main-thread singleton) here because
    _merge_paddle_with_ppstructure() is always called from _ocr_attempt()
    which runs in a single thread per image. Thread-local engine not needed.
    """
    paddle_words: list = []
    pp_result:    list = []

    try:
        raw          = _get_paddle_ocr().ocr(img_array, cls=True)
        paddle_words = _paddle_words_to_flat(raw)
        logger.info(
            "PaddleOCR produced %d words for image merge", len(paddle_words)
        )
    except Exception:
        logger.warning("PaddleOCR failed in image merge", exc_info=True)

    try:
        pp_result = _get_pp_structure()(img_array)
    except Exception:
        logger.warning(
            "PP-Structure failed in image merge — returning PaddleOCR flat text",
            exc_info=True
        )
        return "\n".join(w["text"] for w in paddle_words)

    if not pp_result and not paddle_words:
        return ""

    if not pp_result:
        return "\n".join(w["text"] for w in paddle_words)

    words_in_tables: set  = set()
    table_blocks:    list = []

    for region in pp_result:
        region_type = region.get("type", "").lower()
        res         = region.get("res", None)
        region_bbox = region.get("bbox", None)

        if region_type == "table" and isinstance(res, dict):
            html = res.get("html", "")
            if not html:
                continue

            if paddle_words and region_bbox:
                rows = _extract_table_with_plumber_text(
                    html, paddle_words, region_bbox,
                    scale_x=1.0, scale_y=1.0,
                    words_in_tables=words_in_tables,
                )
                if rows:
                    table_blocks.append(
                        "[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]"
                    )
            else:
                table_text = _html_table_to_text(html)
                if table_text.strip():
                    table_blocks.append(f"[TABLE]\n{table_text}\n[/TABLE]")

    parts: list = []

    if paddle_words:
        remaining = [
            w for i, w in enumerate(paddle_words)
            if i not in words_in_tables
        ]
        if remaining:
            remaining.sort(
                key=lambda w: (round(w["top"] / 10) * 10, w["x0"])
            )
            non_table = " ".join(w["text"] for w in remaining)
            if non_table.strip():
                parts.append(non_table)

    parts.extend(table_blocks)
    return "\n\n".join(parts)


# ══════════════════════════════════════════════════════════════════════════════
#  PARSERS  (unchanged logic — SP-1 engine fix applied to _merge call)
# ══════════════════════════════════════════════════════════════════════════════

def _ocr_attempt(
    file_path:   str,
    progress_cb  = None,
    strategy:    dict | None = None,
    attempt_num: int = 1,
) -> str:
    import numpy as np
    from PIL import Image

    def _cb(pct, msg):
        if progress_cb:
            prefix = f"[Attempt {attempt_num}] " if attempt_num > 1 else ""
            progress_cb(pct, prefix + msg)

    preprocess_kwargs = {
        k: v for k, v in (strategy or {}).items()
        if k in ("sharpness", "contrast", "brightness", "dpi_scale")
    }

    _cb(5, "Stage 1: Preprocessing image…")
    tmp_color, tmp_gray = _preprocess_image(file_path, **preprocess_kwargs)

    try:
        _cb(20, "Stage 2: Running PaddleOCR + PP-Structure merge…")
        ocr_text = ""
        arr      = None
        try:
            arr      = np.array(Image.open(tmp_gray).convert("RGB"))
            ocr_text = _merge_paddle_with_ppstructure(arr)
            logger.info(
                "Image merge produced %d chars for %s",
                len(ocr_text), Path(file_path).name
            )
        except Exception:
            logger.warning(
                "Image merge failed — falling back to plain PaddleOCR",
                exc_info=True
            )
            try:
                if arr is None:
                    arr = np.array(Image.open(tmp_gray).convert("RGB"))
                res = _get_paddle_ocr().ocr(arr, cls=True)
                if res and res[0]:
                    for line in res[0]:
                        if line[1][1] > 0.5:
                            ocr_text += line[1][0] + "\n"
                logger.info(
                    "Plain PaddleOCR fallback produced %d chars", len(ocr_text)
                )
            except Exception:
                logger.warning(
                    "Plain PaddleOCR also failed — continuing with VLM only",
                    exc_info=True
                )
                ocr_text = ""

        _cb(35, "Stage 2b: Checking extraction quality…")

        quality    = _text_quality_score(ocr_text)
        char_count = len(ocr_text.strip())

        if char_count >= 500 and quality >= 0.75:
            logger.info(
                "Image merge quality sufficient (%d chars, %.0f%% readable) — "
                "skipping Gemini VLM. Saving 1 API call.",
                char_count, quality * 100
            )
            _cb(100, "Done (image merge — VLM skipped).")
            return ocr_text

        logger.info(
            "Image merge quality insufficient (%d chars, %.0f%% readable) — "
            "proceeding to Gemini VLM.",
            char_count, quality * 100
        )

        _cb(40, "Stage 3: Sending to Gemini VLM…")

        import PIL.Image as _PILImage
        from google.genai import types as _gtypes

        ref = (
            f"\nOCR REFERENCE (PP-StructureV3 pre-read):\n{ocr_text}\n"
            if ocr_text.strip() else ""
        )

        vlm_prompt = (
            "You are an expert document transcription assistant for a bank.\n"
            "This is a scanned loan-related form or document.\n\n"
            "INSTRUCTIONS:\n"
            "1. Transcribe ALL visible text exactly as it appears.\n"
            "2. Format each field as:  Field Label: value\n"
            "3. Empty fields → [EMPTY]\n"
            "4. Checked checkboxes → [CHECKED], unchecked → [UNCHECKED]\n"
            "5. If text is illegible or partially visible → [UNCLEAR]\n"
            "6. Preserve table structure: use  |  to separate columns.\n"
            "7. Handle multi-column layouts: read left column fully, then right.\n"
            "8. Read order: left→right, top→bottom.\n"
            "9. Include ALL numbers, amounts, dates, ID numbers exactly.\n"
            "10. Do NOT guess or fabricate values. If unsure, write [UNCLEAR].\n"
            + ref
        )

        _VLM_MODELS = ["gemini-2.5-flash", "gemini-2.5-flash-lite"]

        client  = _get_gemini()
        pil_img = _PILImage.open(tmp_color).convert("RGB")

        resp = None
        for _vlm_model in _VLM_MODELS:
            try:
                resp = client.models.generate_content(
                    model    = _vlm_model,
                    contents = [vlm_prompt, pil_img],
                    config   = _gtypes.GenerateContentConfig(
                        max_output_tokens=4000),
                )
                _cb(90, f"Stage 3: VLM complete ({_vlm_model})…")
                break
            except Exception as _e:
                if ("429" in str(_e)
                        or "quota" in str(_e).lower()
                        or "resource_exhausted" in str(_e).lower()):
                    logger.warning(
                        "VLM model %s quota hit — trying next model", _vlm_model
                    )
                    continue
                raise

        if resp is None:
            if ocr_text.strip():
                logger.warning(
                    "VLM quota exhausted — returning PP-Structure result "
                    "(%d chars). May be incomplete.",
                    char_count
                )
                return (
                    f"⚠ VLM quota exhausted — using PP-Structure result only "
                    f"(may be incomplete for low-quality scans):\n\n{ocr_text}"
                )
            return "[VLM quota exhausted on all models — please try again later]"

        text = ""
        try:
            text = resp.text
        except Exception:
            pass
        if not text or not text.strip():
            try:
                text = "".join(
                    p.text for p in resp.candidates[0].content.parts
                    if hasattr(p, "text") and p.text
                )
            except Exception:
                pass
        if not text or not text.strip():
            finish = ""
            try:
                finish = str(resp.candidates[0].finish_reason)
            except Exception:
                pass
            if ocr_text.strip():
                logger.warning(
                    "VLM returned no text (finish_reason: %s) — "
                    "returning PP-Structure result (%d chars).",
                    finish or "unknown", char_count
                )
                return ocr_text
            text = (
                f"[VLM returned no text — finish_reason: {finish or 'unknown'}]"
            )

        _cb(100, "Done.")
        return text

    finally:
        for tmp in (tmp_color, tmp_gray):
            if tmp and tmp != file_path and os.path.exists(tmp):
                try:
                    os.remove(tmp)
                except OSError:
                    logger.warning("Could not delete temp file: %s", tmp)


def _parse_image(file_path: str, progress_cb=None) -> str:
    text = _ocr_attempt(file_path, progress_cb, strategy=None, attempt_num=1)

    clean = text.strip()
    is_error_response = clean.startswith("[VLM") or clean.startswith("⚠")
    response_too_short = len(clean) < _MIN_RESPONSE_CHARS and not is_error_response

    if not response_too_short:
        logger.info(
            "OCR attempt 1 produced %d chars — looks complete, no retry needed.",
            len(clean)
        )
        return text

    logger.warning(
        "OCR attempt 1 response only %d chars — retrying with High Contrast.",
        len(clean)
    )
    if progress_cb:
        progress_cb(
            10,
            f"Short response ({len(clean)} chars) — retrying with High Contrast…"
        )

    strategy   = _RETRY_STRATEGIES[0]
    retry_text = _ocr_attempt(
        file_path, progress_cb,
        strategy    = strategy,
        attempt_num = 2,
    )

    best = retry_text if len(retry_text.strip()) > len(clean) else text
    if len(retry_text.strip()) <= len(clean):
        logger.info(
            "Retry did not improve result (%d vs %d chars) — keeping attempt 1.",
            len(retry_text.strip()), len(clean)
        )
    else:
        logger.info(
            "Retry improved result: %d → %d chars.",
            len(clean), len(retry_text.strip())
        )

    if len(best.strip()) < _MIN_RESPONSE_CHARS:
        warn = (
            "\n⚠ Short response after 2 attempt(s). "
            "This document may need manual review.\n\n"
        )
        logger.warning(
            "Both OCR attempts produced short responses — manual review recommended."
        )
        return warn + best

    return best


def _parse_pdf(file_path: str, progress_cb=None) -> str:
    import pdfplumber

    _TEXT_DENSITY_THRESHOLD = 150
    _MIN_MEANINGFUL_CHARS   = 200

    total     = 0
    text      = ""
    af        = ""
    avg_chars = 0

    try:
        with pdfplumber.open(file_path) as pdf:
            total = len(pdf.pages)
            if progress_cb:
                progress_cb(5, f"PDF: {total} page(s)…")

            for p in pdf.pages:
                text += p.extract_text() or ""
            for p in pdf.pages:
                if p.annots:
                    for a in p.annots:
                        if (a.get("subtype") == "Widget"
                                and a.get("V") and a.get("T")):
                            af += f"{a['T']}: {a['V']}\n"

            avg_chars = len(text) / max(total, 1)

    except Exception as e:
        logger.warning("pdfplumber failed to open PDF: %s", e, exc_info=True)
        text      = ""
        avg_chars = 0

    if avg_chars >= _TEXT_DENSITY_THRESHOLD:
        word_count   = len(text.split())
        unique_words = len(set(text.lower().split()))
        looks_real   = (
            word_count >= 15
            and unique_words >= 8
            and avg_chars >= _MIN_MEANINGFUL_CHARS
        )
        if looks_real:
            logger.info(
                "PDF text layer is real (%.0f chars/page, %d words, %d unique) — "
                "running two-pass pdfplumber + PP-Structure merge (SP-1 parallel).",
                avg_chars, word_count, unique_words
            )
            if progress_cb:
                progress_cb(15, "Merging pdfplumber + PP-Structure (parallel pages)...")
            try:
                dpi    = 300 if total <= 3 else 250 if total <= 10 else 200
                merged = _merge_pdfplumber_with_ppstructure(
                    file_path, dpi=dpi, progress_cb=progress_cb
                )
                if merged.strip():
                    logger.info(
                        "Two-pass merge produced %d chars — 0 API calls.",
                        len(merged)
                    )
                    if progress_cb:
                        progress_cb(100, "Done (pdfplumber + PP-Structure merge).")
                    if af:
                        return merged + "\n=== FORM FIELDS ===\n" + af
                    return merged
                else:
                    logger.warning(
                        "Two-pass merge returned empty — "
                        "falling back to raw pdfplumber text."
                    )
            except Exception as e:
                logger.warning(
                    "Two-pass merge failed (%s) — "
                    "falling back to raw pdfplumber text.", e
                )
            if af:           return text + "\n=== FORM FIELDS ===\n" + af
            if text.strip(): return text
        else:
            logger.info(
                "PDF has %.0f chars/page but only %d words / %d unique — "
                "scanned PDF with thin text layer detected, falling back to OCR.",
                avg_chars, word_count, unique_words
            )

    if avg_chars > 0:
        logger.info(
            "PDF text layer sparse (%.0f chars/page < %d) — "
            "falling back to image OCR.",
            avg_chars, _TEXT_DENSITY_THRESHOLD
        )

    dpi  = 300 if total <= 3 else 250 if total <= 10 else 200
    text = ""

    try:
        from pdf2image import convert_from_path
        kw = ({"poppler_path": POPPLER_PATH}
              if sys.platform == "win32" and Path(POPPLER_PATH).exists() else {})
        imgs = convert_from_path(
            file_path, dpi=dpi, fmt="jpeg",
            jpegopt={"quality": 95, "optimize": True},
            thread_count=min(4, max(total, 1)), **kw
        )
        for i, img in enumerate(imgs):
            if progress_cb:
                progress_cb(
                    int(10 + (i / max(total, 1)) * 80),
                    f"Page {i+1}/{total}…"
                )
            fd, tp = tempfile.mkstemp(suffix=".jpg")
            os.close(fd)
            try:
                img.save(tp, "JPEG", quality=95)
                text += f"\n=== PAGE {i+1} ===\n{_parse_image(tp)}\n"
            finally:
                if os.path.exists(tp):
                    os.remove(tp)
            del img

    except ImportError:
        text += "\n[pdf2image not installed — pip install pdf2image]"
    except Exception as e:
        logger.exception("PDF image fallback failed")
        text += f"\n[PDF render error: {e}]"

    return text


def _parse_excel(file_path: str, progress_cb=None) -> str:
    if progress_cb:
        progress_cb(10, "Reading Excel…")
    text = ""
    try:
        import openpyxl
        from datetime import datetime

        wb = openpyxl.load_workbook(file_path, data_only=True)

        def _get_merged_map(ws):
            merged = {}
            for rng in ws.merged_cells.ranges:
                label = (f"[MERGED {rng.min_row},{rng.min_col}"
                         f"→{rng.max_row},{rng.max_col}]")
                for r in range(rng.min_row, rng.max_row + 1):
                    for c in range(rng.min_col, rng.max_col + 1):
                        merged[(r, c)] = label
            return merged

        def _format_cell(cell):
            val = cell.value
            if val is None:
                return ""
            nf = cell.number_format or ""
            if isinstance(val, datetime):
                return (val.strftime("%Y-%m-%d %H:%M")
                        if val.hour or val.minute
                        else val.strftime("%Y-%m-%d"))
            if isinstance(val, (int, float)):
                if "%" in nf:
                    return f"{val * 100:.2f}%"
                if any(c in nf for c in ("$", "₱", "€", "£", "¥")):
                    sym = next(
                        (c for c in ("$", "₱", "€", "£", "¥") if c in nf), "$"
                    )
                    return f"{sym}{val:,.2f}"
                if "," in nf and "." in nf:
                    return f"{val:,.2f}"
                if isinstance(val, float) and val == int(val):
                    return str(int(val))
            return str(val)

        for sn in wb.sheetnames:
            ws = wb[sn]
            if ws.sheet_state != "visible":
                text += f"\n=== {sn} [HIDDEN — SKIPPED] ===\n"
                continue
            text += f"\n=== SHEET: {sn} ===\n"
            merged_map = _get_merged_map(ws)
            if merged_map:
                text += f"  [Merged regions: {len(set(merged_map.values()))}]\n"
            for row in ws.iter_rows():
                cells_out = []
                for cell in row:
                    coord = (cell.row, cell.column)
                    if coord in merged_map:
                        tag      = merged_map[coord]
                        origin_r = int(tag.split(",")[0].split("[MERGED ")[1])
                        origin_c = int(tag.split("→")[0].split(",")[1])
                        if cell.row == origin_r and cell.column == origin_c:
                            cells_out.append(_format_cell(cell))
                    else:
                        cells_out.append(_format_cell(cell))
                row_str = " | ".join(c for c in cells_out if c)
                if row_str.strip():
                    text += row_str + "\n"

        if progress_cb:
            progress_cb(95, "Done…")
        return text or "[Empty workbook]"

    except Exception as primary_err:
        if progress_cb:
            progress_cb(50, "openpyxl failed, trying pandas…")
        try:
            import pandas as pd
            text = f"[openpyxl error: {primary_err} — falling back to pandas]\n\n"
            xf   = pd.ExcelFile(file_path)
            for sn in xf.sheet_names:
                df = pd.read_excel(xf, sheet_name=sn, header=None)
                text += f"\n=== SHEET: {sn} ===\n"
                if df.empty:
                    text += "[Empty]\n"
                    continue
                for _, row in df.iterrows():
                    parts = [str(v).strip() for v in row
                             if str(v).strip() not in ("", "nan")]
                    if parts:
                        text += " | ".join(parts) + "\n"
            if progress_cb:
                progress_cb(95, "Done (pandas fallback)…")
            return text or "[Empty workbook]"
        except Exception as fallback_err:
            return (f"[Excel read failed]\n"
                    f"openpyxl: {primary_err}\npandas: {fallback_err}")


def _parse_docx(file_path: str, progress_cb=None) -> str:
    import docx
    if progress_cb:
        progress_cb(10, "Reading Word…")
    doc   = docx.Document(file_path)
    text  = ""
    total = len(doc.paragraphs)
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            text += p.text + "\n"
        if progress_cb and i % 10 == 0:
            progress_cb(int(10 + (i / max(total, 1)) * 70), f"Para {i+1}…")
    for t in doc.tables:
        for r in t.rows:
            rt = " | ".join(c.text.strip() for c in r.cells if c.text.strip())
            if rt:
                text += rt + "\n"
    if progress_cb:
        progress_cb(95, "Done…")
    return text or "[No text]"


def _parse_csv(file_path: str, progress_cb=None) -> str:
    import pandas as pd
    if progress_cb:
        progress_cb(10, "Reading CSV…")
    try:
        df = pd.read_csv(file_path)
    except Exception as e:
        return f"[CSV error: {e}]"
    if df.empty:
        return "[Empty CSV]"
    text = ""
    if len(df.columns) == 2:
        for _, r in df.iterrows():
            fld = str(r.iloc[0]).strip()
            val = str(r.iloc[1]).strip()
            if fld and fld.lower() != "nan":
                text += f"{fld}: {val if val.lower() != 'nan' else '[EMPTY]'}\n"
    else:
        text += "=== CSV ===\n" + "|".join(df.columns) + "\n\n"
        for i, r in df.iterrows():
            text += f"--- {i+1} ---\n"
            for c in df.columns:
                text += f"{c}: {str(r[c]).strip()}\n"
            text += "\n"
    if progress_cb:
        progress_cb(95, "Done…")
    return text or "[No data]"


def _parse_text(file_path: str, progress_cb=None) -> str:
    if progress_cb:
        progress_cb(20, "Reading…")
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            t = Path(file_path).read_text(encoding=enc).strip()
            if progress_cb:
                progress_cb(90, "Done…")
            return t
        except (UnicodeDecodeError, LookupError):
            continue
    return "[Encoding error]"


# ══════════════════════════════════════════════════════════════════════════════
#  PUBLIC ENTRY POINTS
# ══════════════════════════════════════════════════════════════════════════════

def extract(file_path: str, progress_cb=None) -> str:
    ext = Path(file_path).suffix.lower()
    try:
        if ext in IMAGE_EXTS:
            result = _parse_image(file_path, progress_cb)
        elif ext == ".pdf":
            result = _parse_pdf(file_path, progress_cb)
        elif ext in (".docx", ".doc"):
            result = _parse_docx(file_path, progress_cb)
        elif ext in (".xlsx", ".xls"):
            result = _parse_excel(file_path, progress_cb)
        elif ext == ".csv":
            result = _parse_csv(file_path, progress_cb)
        else:
            result = _parse_text(file_path, progress_cb)

        if is_cic_report(result):
            result = format_cic_for_analysis(result)

        return result

    except EnvironmentError as e:
        return f"⚠ Config error:\n{e}"
    except ImportError as e:
        pkg = str(e).split("'")[1] if "'" in str(e) else str(e)
        return f"⚠ Missing package: {pkg}\n\npip install {pkg}"
    except Exception as e:
        logger.exception("Extraction failed for %s", file_path)
        return f"⚠ Error:\n{type(e).__name__}: {e}"


# ── File role detection (unchanged) ──────────────────────────────────────────

_LOAN_APP_SHEETS = {
    "cibi", "ci/bi", "cashflow", "cash flow",
    "approval form", "credit scoring", "worksheet",
    "proposed approval form",
}

_RUBRIC_SHEETS = {
    "template-business", "template salary", "template business",
    "parameters-business", "salary-criteria", "criteria-business",
    "users guide", "condition if salary",
}


def _detect_file_role(file_path: str) -> str:
    from pathlib import Path
    import re

    ext  = Path(file_path).suffix.lower()
    name = Path(file_path).stem.lower()

    cic_name_hints = ("cic", "credit_report", "credit report",
                      "creditreport", "cic_report")
    if any(h in name for h in cic_name_hints):
        return "cic_report"

    if ext in (".xlsx", ".xls"):
        rubric_hints = ("rubric", "scoring", "criteria", "matrix",
                        "score_sheet", "scoresheet")
        if any(h in name for h in rubric_hints):
            return "scoring_rubric"
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            ws = wb.active
            header_text = " ".join(
                str(ws.cell(r, c).value or "")
                for r in range(1, 4)
                for c in range(1, 8)
            ).lower()
            wb.close()
            if any(k in header_text for k in ("rubric","scoring","criteria","score")):
                return "scoring_rubric"
        except Exception:
            pass
        return "unknown"

    snippet = ""
    try:
        if ext == ".pdf":
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                pages = pdf.pages[:3]
                snippet = " ".join(p.extract_text() or "" for p in pages)
        elif ext in (".txt", ".csv", ".md"):
            with open(file_path, encoding="utf-8", errors="replace") as fh:
                snippet = fh.read(4000)
        elif ext == ".docx":
            import docx
            doc = docx.Document(file_path)
            snippet = " ".join(p.text for p in doc.paragraphs[:40])
    except Exception:
        snippet = name

    snippet_l = snippet.lower()

    if is_cic_report(snippet):
        return "cic_report"

    worksheet_kw = ("sources of income", "household expenses",
                    "payslip", "salary worksheet", "income worksheet",
                    "average monthly income", "net disposable")
    if sum(1 for kw in worksheet_kw if kw in snippet_l) >= 2:
        return "worksheet"

    finstat_kw = ("balance sheet", "income statement", "cash flow statement",
                  "profit and loss", "stockholders equity", "retained earnings")
    if sum(1 for kw in finstat_kw if kw in snippet_l) >= 2:
        return "financial_stmt"

    loan_kw = ("loan application", "credit investigation", "background investigation",
               "ci/bi", "ci-bi", "cashflow", "loan purpose", "loan amount requested",
               "credit approval", "applicant name", "date of application")
    if sum(1 for kw in loan_kw if kw in snippet_l) >= 2:
        return "loan_application"

    support_kw = ("certificate of employment", "payslip", "pay slip",
                  "government id", "passport", "birth certificate",
                  "barangay clearance", "police clearance")
    if any(kw in snippet_l for kw in support_kw):
        return "supporting_doc"

    return "unknown"


def extract_multiple(file_paths: list[str], progress_cb=None) -> str:
    if not file_paths:
        return "[No files provided]"

    total   = len(file_paths)
    results = []

    for i, fp in enumerate(file_paths):
        name = Path(fp).name
        ext  = Path(fp).suffix.lower()

        base_pct = int(5 + (i / total) * 90)
        end_pct  = int(5 + ((i + 1) / total) * 90)

        def _scoped_cb(pct, stage, _base=base_pct, _end=end_pct):
            if progress_cb:
                mapped = int(_base + (pct / 100) * (_end - _base))
                progress_cb(mapped, f"[{i+1}/{total}] {stage}")

        if progress_cb:
            progress_cb(base_pct, f"Reading file {i+1}/{total}: {name}…")

        if ext in (".xlsx", ".xls"):
            role = _detect_file_role(fp)
            if role == "loan_application":
                label = f"LOAN APPLICATION FILE: {name}"
            elif role == "scoring_rubric":
                label = f"SCORING RUBRIC / TEMPLATE FILE: {name}"
            else:
                label = f"EXCEL FILE: {name}"
        else:
            label = f"FILE: {name}"

        text = extract(fp, _scoped_cb)

        divider = "═" * 60
        results.append(
            f"{divider}\n"
            f"{label}\n"
            f"{divider}\n"
            f"{text}\n"
        )

    if progress_cb:
        progress_cb(98, "Merging files…")

    combined = (
        f"=== COMBINED DOCUMENT PACKAGE ({total} file(s)) ===\n\n"
        + "\n".join(results)
    )

    if progress_cb:
        progress_cb(100, "Done.")

    return combined

# ══════════════════════════════════════════════════════════════════════════════
#  Data classes  (identical to the originals — no breaking change)
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class BankCIRow:
    subject:       str  = ""
    remarks_raw:   str  = ""   # canonical form after fuzzy matching
    signature_raw: str  = ""   # raw VLM output for the signature column
    has_signature: bool = False
    row_verdict:   str  = ""   # POSITIVE | NEGATIVE | INCONCLUSIVE
    row_reason:    str  = ""


@dataclass
class BankCIResult:
    verdict:    str             = "UNCERTAIN"
    proceed:    bool            = False
    summary:    str             = ""
    details:    str             = ""
    full_text:  str             = ""
    rows:       list[BankCIRow] = field(default_factory=list)
    raw_paddle: str             = ""   # kept for API compat; holds raw VLM text
    error:      Optional[str]   = None


# ══════════════════════════════════════════════════════════════════════════════
#  Canonical remarks reference
# ══════════════════════════════════════════════════════════════════════════════

# Each entry: (canonical_label, is_adverse, list_of_known_variants_lowercase)
_REMARKS_CANON: list[tuple[str, bool, list[str]]] = [
    # ── Adverse ───────────────────────────────────────────────────────────
    ("PAST DUE",     True,  ["past due", "pastdue", "past-due", "past oue",
                              "pust due", "past do", "past dve", "p4st due",
                              "past d", "pastd"]),
    ("NPL",          True,  ["npl", "n.p.l", "non-performing", "non performing",
                              "mpl", "np1", "npi", "npl.", "npl-"]),
    ("ITL",          True,  ["itl", "i.t.l", "in the litigation", "litigation",
                              "iti", "1tl", "it1", "itl.", "itl-", "itl,"]),
    ("DELINQUENT",   True,  ["delinquent", "delinq", "delinquint"]),
    ("DEFAULT",      True,  ["default", "defaulted"]),
    ("WRITTEN OFF",  True,  ["written off", "write-off", "writeoff",
                              "chargeoff", "charge-off", "charged off"]),
    ("DISHONORED",   True,  ["dishonored", "dishonoured", "dshonored"]),
    ("BOUNCED",      True,  ["bounced", "bounce", "nsf", "insufficient funds",
                              "insuf funds"]),
    ("OVERDRAFT",    True,  ["overdraft", "overdrawn", "od"]),
    ("BLACKLISTED",  True,  ["blacklisted", "blacklist", "black listed"]),
    ("UNPAID",       True,  ["unpaid", "un-paid"]),
    ("OVERDUE",      True,  ["overdue", "over due", "over-due"]),
    ("RESTRUCTURED", True,  ["restructured", "restructure"]),
    # ── Clean ─────────────────────────────────────────────────────────────
    ("NCD",          False, ["ncd", "n.c.d", "ncd.", "ncd-", "nco", "ncb",
                              "nco.", "no contrary data", "no contrary",
                              "no derogatory", "no adverse", "no adverse record",
                              "no contrary finding", "no findings", "no record of",
                              "nocontrarydata", "noderogatory", "noadverse"]),
    ("CURRENT",      False, ["current", "currnet", "curent", "currant",
                              "currrent", "cuurent", "curren", "currnt",
                              "curent", "currt"]),
    ("CLEAN",        False, ["clean", "clinet", "cleant", "cleen", "cleam",
                              "clcan", "clent", "clen", "clien", "claen",
                              "claan", "cleon", "cleann", "clenn", "clearn"]),
    ("PAID",         False, ["paid", "payed", "fully paid", "full paid"]),
    ("CLOSED",       False, ["closed", "close", "clsd"]),
    ("SETTLED",      False, ["settled", "settld", "setled"]),
    ("GOOD STANDING",False, ["good standing", "good std", "goodstanding"]),
    ("NO RECORD",    False, ["no record", "norecord", "no records"]),
]

# Flat map: lowercase variant → (canonical, is_adverse)
_VARIANT_MAP: dict[str, tuple[str, bool]] = {}
for _canon, _adverse, _variants in _REMARKS_CANON:
    for _v in _variants:
        _VARIANT_MAP[_v.lower()] = (_canon, _adverse)
    # also map the canonical itself
    _VARIANT_MAP[_canon.lower()] = (_canon, _adverse)


def _levenshtein(a: str, b: str) -> int:
    """Standard Levenshtein distance."""
    if len(a) < len(b):
        a, b = b, a
    if not b:
        return len(a)
    prev = list(range(len(b) + 1))
    for ca in a:
        curr = [prev[0] + 1]
        for j, cb in enumerate(b, 1):
            curr.append(min(prev[j] + 1, curr[j - 1] + 1, prev[j - 1] + (ca != cb)))
        prev = curr
    return prev[-1]


def _normalize_remark(raw: str) -> tuple[str, bool]:
    """
    Map a raw VLM remark string to (canonical_label, is_adverse).

    Strategy
    --------
    1. Exact / substring match against all known variants (fast path).
    2. Fuzzy Levenshtein match against short canonical forms when the
       raw string is short (≤ 8 chars) — handles 'NCO', 'mpl', 'iti'.
    3. If nothing matches → return (raw.strip(), False) so unknown
       remarks are treated as neutral / inconclusive rather than hidden.
    """
    if not raw or not raw.strip():
        return ("", False)

    t = raw.strip().lower()
    t_clean = re.sub(r"[^a-z0-9 ]", "", t).strip()

    # 1a. Exact match
    if t in _VARIANT_MAP:
        return _VARIANT_MAP[t]
    if t_clean in _VARIANT_MAP:
        return _VARIANT_MAP[t_clean]

    # 1b. Substring match (e.g. "ncd." → contains "ncd")
    for variant, result in _VARIANT_MAP.items():
        if variant in t or t in variant:
            return result

    # 2. Fuzzy match for short tokens (typos like "NCO" → "NCD")
    if len(t_clean) <= 8:
        best_dist  = 999
        best_canon = None
        best_adv   = False
        for variant, (canon, adv) in _VARIANT_MAP.items():
            if abs(len(variant) - len(t_clean)) > 3:
                continue
            d = _levenshtein(t_clean, variant)
            if d < best_dist:
                best_dist  = d
                best_canon = canon
                best_adv   = adv
        # Accept if edit distance ≤ 2 for very short strings, ≤ 3 for longer
        threshold = 2 if len(t_clean) <= 4 else 3
        if best_dist <= threshold and best_canon:
            logger.info(
                "Fuzzy-matched remark '%s' → '%s' (dist=%d)",
                raw, best_canon, best_dist
            )
            return (best_canon, best_adv)

    # 3. No match found — return as-is, non-adverse
    logger.info("Remark '%s' not matched to canonical list — treating as neutral.", raw)
    return (raw.strip(), False)


# ══════════════════════════════════════════════════════════════════════════════
#  VLM extraction
# ══════════════════════════════════════════════════════════════════════════════

_VLM_PROMPT = """\
You are a credit officer at Banco San Vicente (BSV), a rural bank in the \
Philippines. You are reviewing a Bank CI (Credit Investigation) document.

Your ONLY task is to extract the data table from this document.

OUTPUT FORMAT — respond with ONLY the table rows below, no other text:
BANK: <bank name> | REMARKS: <remarks text> | SIGNATURE: YES or NO

Rules:
1. Extract one row per bank/institution listed in the document.
2. REMARKS: copy the handwritten or typed text exactly as you see it in the \
Remarks column. Do not interpret or correct it — even if it looks like a \
typo (e.g. write "NCO" not "NCD", write "mpl" not "NPL"). Copy verbatim.
3. SIGNATURE: write YES if there is any ink mark, initials, or signature in \
the Signature / Informant column for that row. Write NO otherwise.
4. If a row has no remarks at all, write REMARKS: EMPTY.
5. Skip header rows, address lines, salutation lines, and footnotes.
6. Do not add any explanation, greeting, or summary — only the rows.

Example output:
BANK: BDO Unibank | REMARKS: NCD | SIGNATURE: YES
BANK: Metrobank | REMARKS: past due | SIGNATURE: NO
BANK: PNB | REMARKS: EMPTY | SIGNATURE: YES
"""

_VLM_MODELS = ["gemini-2.5-flash", "gemini-2.5-flash-lite"]


def _call_vlm(pil_img, progress_cb=None) -> str:
    """Send a single PIL image to Gemini VLM and return the raw text response."""
    from google.genai import types as _gtypes

    client = _get_gemini()

    for model in _VLM_MODELS:
        try:
            resp = client.models.generate_content(
                model    = model,
                contents = [_VLM_PROMPT, pil_img.convert("RGB")],
                config   = _gtypes.GenerateContentConfig(
                    max_output_tokens=1024
                ),
            )
            # Extract text robustly
            text = ""
            try:
                text = resp.text or ""
            except Exception:
                pass
            if not text:
                try:
                    text = "".join(
                        p.text for p in resp.candidates[0].content.parts
                        if hasattr(p, "text") and p.text
                    )
                except Exception:
                    pass
            if text.strip():
                logger.info("Bank CI VLM response from %s (%d chars).", model, len(text))
                return text
        except Exception as exc:
            if any(k in str(exc) for k in ("429", "quota", "resource_exhausted")):
                logger.warning("VLM model %s quota hit — trying next.", model)
                continue
            logger.warning("VLM error on model %s: %s", model, exc)
            break   # non-quota error — don't retry on next model

    return ""   # all models failed / quota exhausted


def _parse_vlm_rows(vlm_text: str) -> list[BankCIRow]:
    """
    Parse the structured VLM output into BankCIRow objects.

    Expected line format:
        BANK: <name> | REMARKS: <text> | SIGNATURE: YES or NO

    Returns a list of BankCIRow with:
      - subject       : bank name
      - remarks_raw   : canonical form after fuzzy matching
      - signature_raw : raw VLM output ("YES" / "NO")
      - has_signature : True if VLM said YES
      - row_verdict   : POSITIVE / NEGATIVE / INCONCLUSIVE
      - row_reason    : short explanation
    """
    rows: list[BankCIRow] = []
    _ROW_RE = re.compile(
        r"BANK\s*:\s*(?P<bank>[^|]+)"
        r"\|\s*REMARKS\s*:\s*(?P<remarks>[^|]+)"
        r"\|\s*SIGNATURE\s*:\s*(?P<sig>YES|NO)",
        re.IGNORECASE,
    )

    for line in vlm_text.splitlines():
        line = line.strip()
        if not line:
            continue
        m = _ROW_RE.search(line)
        if not m:
            continue

        bank     = m.group("bank").strip()
        raw_rem  = m.group("remarks").strip()
        raw_sig  = m.group("sig").strip().upper()

        if not bank or bank.upper() in ("BANK", "N/A", "-", "—"):
            continue

        has_sig = raw_sig == "YES"

        # Fuzzy-match the remark to the canonical list
        if raw_rem.upper() in ("EMPTY", "-", "—", "N/A", ""):
            canonical_rem = ""
            is_adverse    = False
        else:
            canonical_rem, is_adverse = _normalize_remark(raw_rem)

        # Determine row verdict
        if is_adverse:
            verdict = "NEGATIVE"
            reason  = f"Adverse remark: '{canonical_rem}' (raw: '{raw_rem}')"
        elif canonical_rem:
            verdict = "POSITIVE"
            reason  = (
                f"Clean remark: '{canonical_rem}'"
                + (" + signature present" if has_sig else "")
            )
        elif has_sig:
            verdict = "POSITIVE"
            reason  = "No remarks; signature present → no adverse finding"
        else:
            verdict = "INCONCLUSIVE"
            reason  = "Empty remarks and no signature detected"

        row = BankCIRow(
            subject       = bank,
            remarks_raw   = canonical_rem or raw_rem,
            signature_raw = raw_sig,
            has_signature = has_sig,
            row_verdict   = verdict,
            row_reason    = reason,
        )
        rows.append(row)

    return rows


# ══════════════════════════════════════════════════════════════════════════════
#  Report builder  (identical output format as before)
# ══════════════════════════════════════════════════════════════════════════════

def _bc_build_full_text(r: BankCIResult) -> str:
    lines = [
        "═" * 62,
        "  BANK CI EXTRACTION REPORT  (BSV — Banco San Vicente)",
        "═" * 62,
        f"  Verdict : {r.verdict}",
        f"  Proceed : {'YES' if r.proceed else 'NO'}",
        f"  Summary : {r.summary}",
        "─" * 62,
    ]
    if r.rows:
        lines.append(
            f"  {'#':<4}{'Bank / Institution':<32}{'Remarks':<14}"
            f"{'Sig':<6}{'Verdict'}"
        )
        lines.append("  " + "─" * 64)
        for i, row in enumerate(r.rows, 1):
            rem_disp = (row.remarks_raw or "—")[:13]
            sig_disp = "Yes" if row.has_signature else "No"
            vrd_icon = {
                "POSITIVE":     "✅",
                "NEGATIVE":     "❌",
                "INCONCLUSIVE": "⚠",
            }.get(row.row_verdict, "•")
            lines.append(
                f"  {i:<4}{(row.subject or '—')[:31]:<32}"
                f"{rem_disp:<14}{sig_disp:<6}{vrd_icon} {row.row_verdict}"
            )
    else:
        lines.append("  (no rows parsed — see raw VLM text in raw_paddle field)")
    lines.append("═" * 62)
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
#  Verdict aggregation  (same logic as before)
# ══════════════════════════════════════════════════════════════════════════════

def _aggregate_verdict(result: BankCIResult) -> BankCIResult:
    rows              = result.rows
    negative_rows     = [r for r in rows if r.row_verdict == "NEGATIVE"]
    positive_rows     = [r for r in rows if r.row_verdict == "POSITIVE"]
    inconclusive_rows = [r for r in rows if r.row_verdict == "INCONCLUSIVE"]

    if negative_rows:
        result.verdict = "BAD"
        result.proceed = False
        result.summary = (
            f"ADVERSE: {len(negative_rows)} negative record(s) detected "
            f"out of {len(rows)} bank(s) checked."
        )
        detail_lines = []
        for r in negative_rows:
            detail_lines.append(f"  ❌  [{r.subject}]: {r.row_reason}")
        for r in positive_rows:
            detail_lines.append(f"  ✅  [{r.subject}]: {r.row_reason}")
        for r in inconclusive_rows:
            detail_lines.append(f"  ⚠   [{r.subject}]: {r.row_reason}")
        result.details = "\n".join(detail_lines)

    elif positive_rows:
        result.verdict = "GOOD"
        result.proceed = True
        result.summary = (
            f"CLEAN: {len(positive_rows)} positive record(s), "
            f"no adverse findings across {len(rows)} bank(s) checked."
        )
        detail_lines = []
        for r in positive_rows:
            detail_lines.append(f"  ✅  [{r.subject}]: {r.row_reason}")
        for r in inconclusive_rows:
            detail_lines.append(f"  ⚠   [{r.subject}]: {r.row_reason}")
        result.details = "\n".join(detail_lines)

    else:
        result.verdict = "UNCERTAIN"
        result.proceed = False
        result.summary = (
            f"INCONCLUSIVE: {len(inconclusive_rows)} row(s) could not be "
            f"fully assessed. Manual review recommended."
        )
        result.details = "\n".join(
            f"  ⚠   [{r.subject}]: {r.row_reason}"
            for r in inconclusive_rows
        )

    return result


# ══════════════════════════════════════════════════════════════════════════════
#  Public API
# ══════════════════════════════════════════════════════════════════════════════
def _bc_load_images(file_path: str) -> list:
    from PIL import Image as _PILImage
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        from pdf2image import convert_from_path
        kw = ({"poppler_path": POPPLER_PATH}
              if sys.platform == "win32" and Path(POPPLER_PATH).exists() else {})
        return convert_from_path(file_path, dpi=300, fmt="jpeg", **kw)
    return [_PILImage.open(file_path).convert("RGB")]

def extract_bank_ci_vlm(file_path: str, progress_cb=None) -> BankCIResult:
    """
    VLM-only Bank CI extraction pipeline.

    Steps
    -----
    1. Load document as PIL images.
    2. Send each page image to Gemini VLM with a structured extraction prompt.
    3. Parse the VLM's table output into BankCIRow objects.
    4. Fuzzy-match each Remarks cell to the canonical BSV remarks list.
    5. Aggregate row verdicts → final GOOD / BAD / UNCERTAIN result.

    Parameters
    ----------
    file_path   : path to the Bank CI file (PDF or image).
    progress_cb : optional (pct: int, msg: str) -> None callback for UI updates.

    Returns
    -------
    BankCIResult — identical structure to the old pipeline.
    """
    def _cb(pct: int, msg: str = ""):
        if progress_cb:
            try:
                progress_cb(pct, msg)
            except Exception:
                pass

    result = BankCIResult()

    # ── Step 1: Load images ───────────────────────────────────────────────
    _cb(5, "Loading Bank CI document…")
    try:
        images = _bc_load_images(file_path)
        _cb(10, f"Loaded {len(images)} page(s).")
    except Exception as exc:
        result.error   = f"Failed to load file: {exc}"
        result.verdict = "UNCERTAIN"
        result.summary = f"Could not load document: {exc}"
        result.full_text = _bc_build_full_text(result)
        return result

    if not images:
        result.error   = "No pages found in document."
        result.verdict = "UNCERTAIN"
        result.summary = "Document appears empty."
        result.full_text = _bc_build_full_text(result)
        return result

    # ── Step 2: VLM extraction (one call per page) ────────────────────────
    all_vlm_text: list[str] = []
    all_rows:     list[BankCIRow] = []

    for page_num, pil_img in enumerate(images):
        _cb(
            int(10 + (page_num / max(len(images), 1)) * 60),
            f"VLM reading page {page_num + 1}/{len(images)}…"
        )

        vlm_text = _call_vlm(pil_img, progress_cb)

        if not vlm_text.strip():
            logger.warning(
                "VLM returned empty response for page %d — "
                "quota exhausted or model error.", page_num + 1
            )
            if page_num == 0 and len(images) == 1:
                # Single page and VLM failed entirely — return error result
                result.error   = "VLM quota exhausted or model error."
                result.verdict = "UNCERTAIN"
                result.summary = (
                    "VLM could not read the document. "
                    "Please try again or check API quota."
                )
                result.full_text = _bc_build_full_text(result)
                return result
            continue

        all_vlm_text.append(vlm_text)

        # ── Step 3 & 4: Parse rows + fuzzy-match remarks ──────────────────
        page_rows = _parse_vlm_rows(vlm_text)
        logger.info(
            "Page %d: VLM returned %d row(s).", page_num + 1, len(page_rows)
        )
        all_rows.extend(page_rows)

    # Store raw VLM text in raw_paddle for audit trail / AI context
    result.raw_paddle = "\n---PAGE BREAK---\n".join(all_vlm_text)
    result.rows       = all_rows

    _cb(75, "Aggregating row verdicts…")

    # ── Step 5: Aggregate ─────────────────────────────────────────────────
    if not all_rows:
        # VLM produced text but no parseable rows
        # Fall back to a keyword scan of the raw VLM text
        _cb(80, "No rows parsed — running fallback keyword scan on VLM text…")
        combined = result.raw_paddle.lower()
        adverse_hits = [
            kw for kw in ("past due", "npl", "itl", "delinquent", "default",
                           "written off", "dishonored", "bounced", "nsf",
                           "overdraft", "blacklisted", "unpaid", "overdue")
            if kw in combined
        ]
        clean_hits = [
            kw for kw in ("ncd", "current", "clean", "paid", "settled",
                           "good standing", "no record", "no adverse")
            if kw in combined
        ]

        if adverse_hits:
            result.verdict = "BAD"
            result.proceed = False
            result.summary = f"Adverse keyword(s) found in VLM text: {', '.join(adverse_hits)}"
        elif clean_hits:
            result.verdict = "GOOD"
            result.proceed = True
            result.summary = f"Clean indicator(s) found in VLM text: {', '.join(clean_hits)}"
        else:
            result.verdict = "UNCERTAIN"
            result.proceed = False
            result.summary = (
                "No structured rows could be parsed from the Bank CI document "
                "and no clear keywords found. Manual review recommended."
            )
        result.details   = "Verdict based on raw VLM keyword scan (no structured rows)."
        result.full_text = _bc_build_full_text(result)
        _cb(100, f"Bank CI complete: {result.verdict} (keyword fallback).")
        return result

    result = _aggregate_verdict(result)

    _cb(90, "Building report…")
    result.full_text = _bc_build_full_text(result)
    _cb(100, f"Bank CI complete: {result.verdict}.")
    return result


def bank_ci_to_ai_context(result: BankCIResult) -> str:
    """
    Convert a BankCIResult into a plain-text string suitable for passing
    to the AI prompt in cibi_populator / _ai_check_stage1().

    Returns full_text if populated, otherwise builds a compact summary.
    """
    if result.full_text and result.full_text.strip():
        return result.full_text

    lines = [
        f"Bank CI Verdict: {result.verdict}",
        f"Proceed: {'YES' if result.proceed else 'NO'}",
        f"Summary: {result.summary}",
    ]
    if result.details:
        lines.append(f"Details:\n{result.details}")
    if result.rows:
        lines.append("Row findings:")
        for i, row in enumerate(result.rows, 1):
            lines.append(
                f"  {i}. [{row.subject}] "
                f"Remarks: '{row.remarks_raw or 'EMPTY'}' | "
                f"Signature: {'Yes' if row.has_signature else 'No'} | "
                f"{row.row_verdict}: {row.row_reason}"
            )
    if result.raw_paddle:
        lines.append(f"\nRaw VLM text:\n{result.raw_paddle}")
    return "\n".join(lines)